from flask import Blueprint, jsonify, request, send_from_directory, current_app, render_template, g
from app.services.email_processor import EmailProcessor
from app.services.export_handler import ExportHandler
from app.services.recruiter_notification_service import RecruiterNotificationService
from app.scheduler import get_scheduler_status, pause_scheduler, resume_scheduler, run_job_manually
from app.models.requirement import Requirement, DepartmentEnum, CompanyEnum, ShiftEnum, JobTypeEnum, PriorityEnum
from app.models.profile import Profile
from app.models.user import User
# Legacy Tracker model - deprecated, use Profile.requirement_id relationship instead
# from app.models.tracker import Tracker
from app.database import db
from app.models.meeting import Meeting
from app.middleware.auth import require_admin
from app.middleware.domain_auth import require_domain_auth, require_jwt_domain_auth, ensure_domain_isolation
from app.services.database_manager import database_manager
from flask_jwt_extended import create_access_token, jwt_required, get_jwt_identity
import os
from datetime import datetime, timedelta
from sqlalchemy.orm import scoped_session
from sqlalchemy import func, and_, or_
import re
import logging
import json

def get_db_session():
    """
    Get the correct database session for the current domain.
    Returns domain-specific session if available, otherwise falls back to global session.
    """
    try:
        # Check if we have a domain-specific session
        if hasattr(g, 'db_session') and g.db_session is not None:
            # Verify it's a valid session object
            if hasattr(g.db_session, 'query'):
                return g.db_session
        
        # Fallback to global session for backward compatibility
        # Get the actual session from Flask-SQLAlchemy
        try:
            # This gets the actual SQLAlchemy session
            session = db.session
            if hasattr(session, 'query'):
                return session
            else:
                current_app.logger.error("db.session does not have query method")
                # Try to create a new session from the engine
                from sqlalchemy.orm import sessionmaker
                Session = sessionmaker(bind=db.engine)
                return Session()
        except Exception as session_error:
            current_app.logger.error(f"Error accessing db.session: {str(session_error)}")
            # Last resort: try to create session from engine
            try:
                from sqlalchemy.orm import sessionmaker
                Session = sessionmaker(bind=db.engine)
                return Session()
            except Exception as engine_error:
                current_app.logger.error(f"Error creating session from engine: {str(engine_error)}")
                raise Exception("Cannot create database session")
        
    except Exception as e:
        # If there's any error, log it and re-raise
        current_app.logger.error(f"Critical error in get_db_session: {str(e)}")
        raise e

def format_enum_for_display(value):
    """Convert database enum values to user-friendly display format"""
    if not value:
        return value
    # Replace underscores with spaces and title case each word
    return value.replace('_', ' ').title()

def _get_assigned_recruiters_for_requirement(requirement_id):
    """Get the list of assigned recruiters for a requirement using Assignment model"""
    try:
        from app.models.assignment import Assignment
        from app.models.user import User
        from flask import g
        
        # Use domain-specific session if available, otherwise fall back to global session
        if hasattr(g, 'db_session') and g.db_session is not None:
            session = g.db_session
        else:
            session = db.session
        
        assignments = session.query(Assignment).filter_by(
            requirement_id=requirement_id,
            is_active=True
        ).all()
        
        recruiters = []
        
        for assignment in assignments:
            # Get user from the same session
            user = session.query(User).filter_by(user_id=assignment.user_id).first()
            if user and user.role.value == 'recruiter':
                recruiters.append(user.username)
        
        return recruiters
    except Exception as e:
        current_app.logger.error(f"Error getting assigned recruiters for requirement {requirement_id}: {str(e)}")
        return []

api_bp = Blueprint('api', __name__, url_prefix='/api')



@api_bp.route('/', methods=['GET'])
def api_root():
    """API root endpoint with available endpoints"""
    return jsonify({
        'status': 'success',
        'message': 'Email Filter API',
        'version': '1.0',
        'endpoints': {
            'test_token': '/test-token',
            'fetch_all_emails': '/emails/all',
            'fetch_recruiter_emails': '/emails/recruiter',
            'export_data': '/emails/export',
            'serve_attachment': '/attachments/<filename>',
            'serve_export': '/exports/<filename>',
            'get_requirements': '/requirements',
            'get_requirement': '/requirements/<id>',
            'get_profiles': '/profiles',
            'get_profile': '/profiles/<student_id>',
            'get_tracker_requirements': '/tracker',
            'get_tracker_requirement': '/tracker/<request_id>',
            'update_tracker_requirement': '/tracker/<request_id> (PUT)',
            'delete_tracker_requirement': '/tracker/<request_id> (DELETE)',
            'get_tracker_stats': '/tracker/stats',
            'convert_docx_to_html': '/convert-docx-to-html',
            'download_file': '/download-file',
            'view_file': '/view-file'
        }
    })

@api_bp.route('/test-token', methods=['GET'])
def test_token():
    """Test the Microsoft Graph API token"""
    email_processor = EmailProcessor()
    try:
        result = email_processor.test_token()
        return jsonify(result)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@api_bp.route('/emails/all', methods=['GET'])
@require_domain_auth
def get_all_emails():
    """Get all processed emails"""
    try:
        days = request.args.get('days', default=None, type=int)
        email_processor = EmailProcessor()
        emails = email_processor.fetch_emails(days if days is not None else None)
        
        if not emails:
            return jsonify([])  # Return empty array instead of 404
        
        # Process each email to include clean body and other details
        processed_emails = []
        
        for email in emails:
            if not email:  # Skip None values
                continue
                
            try:
                # Get body content and type
                body = email.get('body', {})
                if isinstance(body, dict):
                    body_content = body.get('content', '')
                    content_type = body.get('contentType', 'text')
                else:
                    body_content = str(body)
                    content_type = 'text'

                # Get sender info
                sender = email.get('from', {})
                if isinstance(sender, dict):
                    sender_address = sender.get('emailAddress', {}).get('address', '')
                else:
                    sender_address = str(sender)

                # Create processed email object
                processed_email = {
                    'id': email.get('id', ''),
                    'subject': email.get('subject', ''),
                    'sender': sender_address,
                    'receivedDateTime': email.get('receivedDateTime', ''),
                    'body': body_content,
                    'body_content_type': content_type,
                    'clean_body': email_processor._clean_text(body_content) if body_content else '',
                    'full_body': body_content,
                    'body_preview': email.get('bodyPreview', ''),
                    'attachments': email.get('attachments', [])
                }

                # Process for requirements only (NO automatic profile extraction)
                try:
                    # Check if this is an RFH email first
                    is_rfh_email = email_processor._is_rfh_email(processed_email)
                    
                    # Process email with requirement creation enabled ONLY for RFH emails
                    profiles_created = 0  # No automatic profile extraction
                    should_create_requirement = is_rfh_email
                    
                    # Process with requirement creation ONLY for RFH emails
                    if should_create_requirement:
                        result = email_processor._process_email(processed_email, create_requirements=True)
                    else:
                        result = email_processor._process_email(processed_email, create_requirements=False)
                    
                    if result:
                        processed_email['profiles_created'] = 0  # No automatic profile extraction
                        processed_email['requirement_created'] = result.get('requirement_created', False)
                        processed_email['requirements_data'] = result.get('requirements_data', None)
                        processed_email['is_rfh_email'] = is_rfh_email
                        processed_email['has_profiles'] = False  # No automatic profile extraction
                        processed_email['thread_id'] = result.get('thread_id')
                        processed_email['existing_requirement_updated'] = result.get('existing_requirement_updated', False)
                        current_app.logger.info(f"Processed email {email.get('id', '')}: profiles=0 (no auto extraction), requirements={1 if result.get('requirement_created') else 0}, is_rfh={is_rfh_email}, has_profiles=False, thread_id={result.get('thread_id')}, updated={result.get('existing_requirement_updated', False)}")
                    else:
                        current_app.logger.warning(f"No result from processing email {email.get('id', '')}")
                        processed_email['profiles_created'] = 0
                        processed_email['requirement_created'] = False
                        processed_email['requirements_data'] = None
                        processed_email['is_rfh_email'] = is_rfh_email
                        processed_email['has_profiles'] = False
                        processed_email['thread_id'] = None
                        processed_email['existing_requirement_updated'] = False
                except Exception as e:
                    current_app.logger.error(f"Error processing email data: {str(e)}")
                    processed_email['profiles_created'] = 0
                    processed_email['requirement_created'] = False
                    processed_email['requirements_data'] = None
                    processed_email['is_rfh_email'] = False
                    processed_email['has_profiles'] = False
                    processed_email['thread_id'] = None
                    processed_email['existing_requirement_updated'] = False

                processed_emails.append(processed_email)
            except Exception as e:
                current_app.logger.error(f"Error processing email {email.get('subject', 'Unknown')}: {str(e)}")
                continue
        
        return jsonify(processed_emails)
    except Exception as e:
        current_app.logger.error(f"Error in get_all_emails: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': f'Failed to fetch emails: {str(e)}'
        }), 500

@api_bp.route('/emails/recruiter', methods=['GET'])
@require_domain_auth
def get_recruiter_emails():
    """Get recruiter emails"""
    try:
        days = request.args.get('days', default=None, type=int)
        email_processor = EmailProcessor()
        emails = email_processor.get_recruiter_emails(days if days is not None else None)
        
        if not emails:
            return jsonify([])  # Return empty array instead of error
        
        # Process each email to include clean body and other details
        processed_emails = []
        for email in emails:
            try:
                # Get body content and type
                body = email.get('body', {})
                if isinstance(body, dict):
                    body_content = body.get('content', '')
                    content_type = body.get('contentType', 'text')
                else:
                    body_content = str(body)
                    content_type = 'text'

                # Get sender info
                sender = email.get('from', {})
                if isinstance(sender, dict):
                    sender_address = sender.get('emailAddress', {}).get('address', '')
                else:
                    sender_address = str(sender)

                # Create processed email object
                processed_email = {
                    'id': email.get('id', ''),
                    'subject': email.get('subject', ''),
                    'sender': sender_address,
                    'receivedDateTime': email.get('receivedDateTime', ''),
                    'body': body_content,
                    'body_content_type': content_type,
                    'clean_body': email_processor._clean_text(body_content) if body_content else '',
                    'full_body': body_content,
                    'body_preview': email.get('bodyPreview', ''),
                    'attachments': email.get('attachments', [])
                }

                # Process for requirements only (NO automatic profile extraction)
                try:
                    # Check if this is an RFH email
                    is_rfh_email = email_processor._is_rfh_email(processed_email)
                    
                    # Process email with requirement creation enabled ONLY for RFH emails
                    if is_rfh_email:
                        result = email_processor._process_email(processed_email, create_requirements=True)
                        if result:
                            processed_email['profiles_created'] = 0  # No automatic profile extraction
                            processed_email['requirement_created'] = result.get('requirement_created', False)
                            processed_email['requirements_data'] = result.get('requirements_data', None)
                            processed_email['is_rfh_email'] = True
                        else:
                            processed_email['profiles_created'] = 0
                            processed_email['requirement_created'] = False
                            processed_email['requirements_data'] = None
                            processed_email['is_rfh_email'] = True
                    else:
                        # For non-RFH emails, no automatic processing
                        processed_email['profiles_created'] = 0
                        processed_email['requirement_created'] = False
                        processed_email['requirements_data'] = None
                        processed_email['is_rfh_email'] = False
                except Exception as e:
                    current_app.logger.error(f"Error processing email data: {str(e)}")
                    processed_email['profiles_created'] = 0
                    processed_email['requirement_created'] = False
                    processed_email['requirements_data'] = None
                    processed_email['is_rfh_email'] = False
                
                processed_emails.append(processed_email)
            except Exception as e:
                current_app.logger.error(f"Error processing email {email.get('subject', 'Unknown')}: {str(e)}")
                continue
        
        return jsonify(processed_emails)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@api_bp.route('/email-refresh-status', methods=['GET'])
def get_email_refresh_status():
    """Get the current email refresh status and last refresh time"""
    try:
        from app.models.system_settings import SystemSettings
        from datetime import datetime
        
        last_refresh_str = SystemSettings.get_setting('last_email_refresh')
        last_refresh = None
        
        if last_refresh_str:
            try:
                last_refresh = datetime.fromisoformat(last_refresh_str)
            except ValueError:
                current_app.logger.warning(f"Invalid last refresh time format: {last_refresh_str}")
                last_refresh = None
        
        return jsonify({
            'last_refresh_time': last_refresh.isoformat() if last_refresh else None,
            'last_refresh_timestamp': last_refresh.timestamp() if last_refresh else None,
            'has_previous_refresh': last_refresh is not None
        })
        
    except Exception as e:
        current_app.logger.error(f"Error getting refresh status: {str(e)}")
        return jsonify({'error': str(e)}), 500

@api_bp.route('/clear-email-refresh-history', methods=['POST'])
def clear_email_refresh_history():
    """Clear the email refresh history (for testing purposes)"""
    try:
        from app.models.system_settings import SystemSettings
        
        # Remove the last refresh time setting
        SystemSettings.set_setting('last_email_refresh', '')
        
        current_app.logger.info("Email refresh history cleared")
        return jsonify({
            'success': True,
            'message': 'Email refresh history cleared successfully'
        })
        
    except Exception as e:
        current_app.logger.error(f"Error clearing refresh history: {str(e)}")
        return jsonify({'error': str(e)}), 500

@api_bp.route('/get-latest-mails', methods=['GET'])
def get_latest_mails():
    """Get latest emails for real-time updates"""
    try:
        from app.models.system_settings import SystemSettings
        from datetime import datetime
        
        # Get last refresh time from database
        last_refresh_str = SystemSettings.get_setting('last_email_refresh')
        last_refresh = None
        
        if last_refresh_str:
            try:
                last_refresh = datetime.fromisoformat(last_refresh_str)
            except ValueError:
                current_app.logger.warning(f"Invalid last refresh time format: {last_refresh_str}")
                last_refresh = None
        
        email_processor = EmailProcessor()
        
        # Update refresh time BEFORE fetching to ensure we don't miss emails
        # that arrive during the fetch process
        current_refresh_time = datetime.utcnow()
        
        # If no last refresh time, use default days (20)
        if not last_refresh:
            days = request.args.get('days', default=20, type=int)
            current_app.logger.info(f"No previous refresh time found, fetching emails for last {days} days")
            emails = email_processor.fetch_emails(days)
            current_app.logger.info(f"Fetched {len(emails) if emails else 0} emails for last {days} days")
        else:
            # Fetch emails only after last refresh time
            current_app.logger.info(f"Fetching emails since last refresh: {last_refresh.isoformat()}")
            emails = email_processor.fetch_emails_since(last_refresh)
            current_app.logger.info(f"Fetched {len(emails) if emails else 0} emails since {last_refresh.isoformat()}")
        
        # Update last refresh time after successful fetch
        try:
            SystemSettings.set_setting('last_email_refresh', current_refresh_time.isoformat())
            current_app.logger.info(f"Updated last refresh time to {current_refresh_time.isoformat()}")
        except Exception as e:
            current_app.logger.error(f"Error updating last refresh time: {str(e)}")
            # Continue processing even if refresh time update fails
        
        if not emails:
            return jsonify([])  # Return empty array instead of 404
        
        # Process each email to include clean body and other details
        processed_emails = []
        
        for email in emails:
            if not email:  # Skip None values
                continue
                
            try:
                # Get body content and type
                body = email.get('body', {})
                if isinstance(body, dict):
                    body_content = body.get('content', '')
                    content_type = body.get('contentType', 'text')
                else:
                    body_content = str(body)
                    content_type = 'text'

                # Get sender info
                sender = email.get('from', {})
                if isinstance(sender, dict):
                    sender_address = sender.get('emailAddress', {}).get('address', '')
                else:
                    sender_address = str(sender)

                # Create processed email object
                processed_email = {
                    'id': email.get('id', ''),
                    'subject': email.get('subject', ''),
                    'sender': sender_address,
                    'receivedDateTime': email.get('receivedDateTime', ''),
                    'body': body_content,
                    'body_content_type': content_type,
                    'clean_body': email_processor._clean_text(body_content) if body_content else '',
                    'full_body': body_content,
                    'body_preview': email.get('bodyPreview', ''),
                    'attachments': email.get('attachments', [])
                }

                # Process for requirements only (NO automatic profile extraction)
                try:
                    # Check if this is an RFH email first
                    is_rfh_email = email_processor._is_rfh_email(processed_email)
                    
                    # Process email with requirement creation enabled ONLY for RFH emails
                    profiles_created = 0  # No automatic profile extraction
                    should_create_requirement = is_rfh_email
                    
                    # Process with requirement creation ONLY for RFH emails
                    if should_create_requirement:
                        result = email_processor._process_email(processed_email, create_requirements=True)
                    else:
                        result = email_processor._process_email(processed_email, create_requirements=False)
                    
                    if result:
                        processed_email['profiles_created'] = 0  # No automatic profile extraction
                        processed_email['requirement_created'] = result.get('requirement_created', False)
                        processed_email['requirements_data'] = result.get('requirements_data', None)
                        processed_email['is_rfh_email'] = is_rfh_email
                        processed_email['has_profiles'] = False  # No automatic profile extraction
                        processed_email['thread_id'] = result.get('thread_id')
                        processed_email['existing_requirement_updated'] = result.get('existing_requirement_updated', False)
                    else:
                        processed_email['profiles_created'] = 0
                        processed_email['requirement_created'] = False
                        processed_email['requirements_data'] = None
                        processed_email['is_rfh_email'] = is_rfh_email
                        processed_email['has_profiles'] = False
                        processed_email['thread_id'] = None
                        processed_email['existing_requirement_updated'] = False
                except Exception as e:
                    current_app.logger.error(f"Error processing email data: {str(e)}")
                    processed_email['profiles_created'] = 0
                    processed_email['requirement_created'] = False
                    processed_email['requirements_data'] = None
                    processed_email['is_rfh_email'] = False
                    processed_email['has_profiles'] = False
                    processed_email['thread_id'] = None
                    processed_email['existing_requirement_updated'] = False

                processed_emails.append(processed_email)
            except Exception as e:
                current_app.logger.error(f"Error processing email {email.get('subject', 'Unknown')}: {str(e)}")
                continue
        
        return jsonify(processed_emails)
    except Exception as e:
        current_app.logger.error(f"Error in get_latest_mails: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': f'Failed to fetch latest emails: {str(e)}'
        }), 500

@api_bp.route('/emails/process', methods=['POST'])
def process_emails():
    """Process only RFH emails and create requirements (no duplicates, one per email)"""
    try:
        email_processor = EmailProcessor()
        all_emails = email_processor.fetch_emails(None)
        created_count = 0
        skipped_invalid = 0
        skipped_duplicate = 0
        skipped_non_rfh = 0
        
        for email in all_emails:
            # Check if this is an RFH email
            processed_email = {
                'id': email.get('id', ''),
                'subject': email.get('subject', ''),
                'sender': email.get('from', {}).get('emailAddress', {}).get('address', ''),
                'receivedDateTime': email.get('receivedDateTime', ''),
                'body': email.get('body', {}).get('content', ''),
                'body_content_type': 'html'
            }
            
            is_rfh_email = email_processor._is_rfh_email(processed_email)
            
            if not is_rfh_email:
                skipped_non_rfh += 1
                continue
            
            # Get body content and type
            body = email.get('body', {})
            if isinstance(body, dict):
                body_content = body.get('content', '')
            else:
                body_content = str(body)

            # Get sender info
            sender = email.get('from', {})
            if isinstance(sender, dict):
                sender_address = sender.get('emailAddress', {}).get('address', '')
            else:
                sender_address = str(sender)

            # Create processed email object
            processed_email = {
                'id': email.get('id', ''),
                'subject': email.get('subject', ''),
                'sender': sender_address,
                'receivedDateTime': email.get('receivedDateTime', ''),
                'body': body_content,
                'body_content_type': 'html'
            }
            
            # Process email with requirement creation enabled (only for RFH emails)
            result = email_processor._process_email(processed_email, create_requirements=True)
            
            if result.get('requirement_created'):
                created_count += 1
            elif result.get('requirements_data') and not result.get('requirement_created'):
                # Valid requirement data but not created (likely duplicate)
                skipped_duplicate += 1
            else:
                # No valid requirement data found
                skipped_invalid += 1
        
        return jsonify({
            'message': f'Processed {len(all_emails)} emails',
            'created': created_count,
            'skipped_invalid': skipped_invalid,
            'skipped_duplicate': skipped_duplicate,
            'skipped_non_rfh': skipped_non_rfh
        })
    except Exception as e:
        current_app.logger.error(f"Error processing emails: {str(e)}")
        return jsonify({'error': str(e)}), 500

@api_bp.route('/profiles/requirement/<string:request_id>', methods=['GET'])
@require_domain_auth
def get_profiles_for_requirement(request_id):
    """Get all profiles for a specific requirement by request_id"""
    try:
        # First check if the requirement exists
        requirement = get_db_session().query(Requirement).filter_by(request_id=request_id).first()
        if not requirement:
            return jsonify({'error': 'Requirement not found'}), 404
        
        # Get all profiles linked to this requirement
        profiles = get_db_session().query(Profile).filter(
            Profile.requirement_id == requirement.requirement_id,
            Profile.deleted_at.is_(None)
        ).all()
        
        # Convert to dictionaries
        profiles_data = []
        for profile in profiles:
            profile_dict = profile.to_dict()
            profiles_data.append(profile_dict)
        
        return jsonify(profiles_data)
        
    except Exception as e:
        current_app.logger.error(f"Error fetching profiles for requirement {request_id}: {str(e)}")
        return jsonify({'error': 'Failed to fetch profiles'}), 500

@api_bp.route('/requirements/<string:request_id>', methods=['GET'])
def get_requirement_by_id(request_id):
    """Get a specific requirement by request_id"""
    try:
        requirement = get_db_session().query(Requirement).filter_by(request_id=request_id).first()
        if not requirement:
            return jsonify({'error': 'Requirement not found'}), 404
        
        return jsonify(requirement.to_dict())
        
    except Exception as e:
        current_app.logger.error(f"Error fetching requirement {request_id}: {str(e)}")
        return jsonify({'error': 'Failed to fetch requirement'}), 500

@api_bp.route('/profiles', methods=['GET'])
@require_domain_auth
def get_profiles():
    """Get all candidate profiles with optional search and filtering"""
    try:
        # Get query parameters
        search = request.args.get('search', '').strip()
        experience_min = request.args.get('experience_min', type=float)
        experience_max = request.args.get('experience_max', type=float)
        ctc_min = request.args.get('ctc_min', type=float)
        ctc_max = request.args.get('ctc_max', type=float)
        location = request.args.get('location', '').strip()
        company = request.args.get('company', '').strip()
        page = request.args.get('page', 1, type=int)
        per_page = request.args.get('per_page', 20, type=int)
        
        # Start with base query, excluding soft-deleted profiles
        query = get_db_session().query(Profile).filter(Profile.deleted_at.is_(None))
        
        # Apply search filter
        if search:
            search_filter = db.or_(
                Profile.candidate_name.ilike(f'%{search}%'),
                Profile.email_id.ilike(f'%{search}%'),
                Profile.contact_no.ilike(f'%{search}%'),
                Profile.key_skills.ilike(f'%{search}%'),
                Profile.current_company.ilike(f'%{search}%'),
                Profile.location.ilike(f'%{search}%')
            )
            query = query.filter(search_filter)
        
        # Apply experience filter
        if experience_min is not None:
            query = query.filter(Profile.total_experience >= experience_min)
        if experience_max is not None:
            query = query.filter(Profile.total_experience <= experience_max)
        
        # Apply CTC filter
        if ctc_min is not None:
            query = query.filter(Profile.ctc_expected >= ctc_min)
        if ctc_max is not None:
            query = query.filter(Profile.ctc_expected <= ctc_max)
        
        # Apply location filter
        if location:
            query = query.filter(Profile.location.ilike(f'%{location}%'))
        
        # Apply company filter
        if company:
            query = query.filter(Profile.current_company.ilike(f'%{company}%'))
        
        # Get total count for pagination
        total_count = query.count()
        
        # Calculate pagination values
        offset = (page - 1) * per_page
        total_pages = (total_count + per_page - 1) // per_page  # Ceiling division
        has_next = page < total_pages
        has_prev = page > 1
        
        # Apply pagination and ordering
        profiles = query.order_by(Profile.created_at.desc()).offset(offset).limit(per_page).all()
        
        return jsonify({
            'profiles': [profile.to_dict() for profile in profiles],
            'pagination': {
                'page': page,
                'per_page': per_page,
                'total': total_count,
                'pages': total_pages,
                'has_next': has_next,
                'has_prev': has_prev
            }
        })
    except Exception as e:
        current_app.logger.error(f"Error fetching profiles data: {str(e)}")
        return jsonify({'error': 'Failed to fetch profiles data'}), 500

@api_bp.route('/profiles/<string:student_id>', methods=['GET'])
def get_profile(student_id):
    """Get a specific profile by student_id"""
    try:
        profile = get_db_session().query(Profile).filter(
            Profile.student_id == student_id,
            Profile.deleted_at.is_(None)
        ).first_or_404()
        return jsonify(profile.to_dict())
    except Exception as e:
        current_app.logger.error(f"Error fetching profile {student_id}: {str(e)}")
        return jsonify({'error': f'Failed to fetch profile {student_id}'}), 500

@api_bp.route('/profiles/<string:student_id>', methods=['DELETE'])
def delete_profile(student_id):
    """Soft delete a profile by student_id"""
    try:
        profile = get_db_session().query(Profile).filter(
            Profile.student_id == student_id,
            Profile.deleted_at.is_(None)
        ).first_or_404()

        # Check if profile is already soft deleted
        if profile.deleted_at:
            return jsonify({
                'success': False,
                'message': f'Profile {student_id} is already deleted'
            }), 400

        # Perform soft delete by setting deleted_at timestamp
        profile.deleted_at = datetime.utcnow()
        get_db_session().commit()

        current_app.logger.info(f"Profile {student_id} soft deleted successfully")
        return jsonify({
            'success': True,
            'message': f'Profile {student_id} deleted successfully'
        })

    except Exception as e:
        current_app.logger.error(f"Error deleting profile {student_id}: {str(e)}")
        get_db_session().rollback()
        return jsonify({
            'success': False,
            'message': f'Failed to delete profile: {str(e)}'
        }), 500

@api_bp.route('/attachments/<path:filename>')
def serve_attachment(filename):
    """Serve attachment files"""
    try:
        return send_from_directory(
            os.path.join(current_app.root_path, 'uploads', 'attachments'),
            filename
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 404

@api_bp.route('/exports/<path:filename>')
def serve_export(filename):
    """Serve exported files"""
    try:
        export_dir = os.path.join(current_app.root_path, 'uploads', 'exports')
        file_path = os.path.join(export_dir, filename)
        
        # Debug logging
        current_app.logger.info(f"Serving export file: {filename}")
        current_app.logger.info(f"Export directory: {export_dir}")
        current_app.logger.info(f"Full file path: {file_path}")
        current_app.logger.info(f"File exists: {os.path.exists(file_path)}")
        
        if not os.path.exists(file_path):
            current_app.logger.error(f"File not found: {file_path}")
            return jsonify({'error': 'File not found'}), 404
        
        return send_from_directory(export_dir, filename, as_attachment=True)
    except Exception as e:
        current_app.logger.error(f"Error serving export file {filename}: {str(e)}")
        return jsonify({'error': str(e)}), 404

@api_bp.route('/emails/export', methods=['POST'])
def export_data():
    """Export requirements data"""
    try:
        data = request.get_json()
        emails = data.get('emails', [])
        export_handler = ExportHandler()
        file_path = export_handler.export_data(emails)
        
        return jsonify({
            'status': 'success',
            'message': 'Data exported successfully',
            'file_path': file_path
        })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

@api_bp.route('/requirements', methods=['GET'])
@require_domain_auth
def get_requirements():
    """Get all job requirements (excluding reply/forward emails)"""
    try:
        current_app.logger.info("Fetching all requirements")
        # Get manual requirements (these should always be included regardless of email subject)
        manual_requirements = get_db_session().query(Requirement).filter(
            Requirement.requirement_id.isnot(None),
            Requirement.is_manual_requirement == True
        ).order_by(Requirement.created_at.desc()).all()
        
        # Get automatic requirements (email-based) with existing filters
        # Join with EmailDetails to filter by email subject prefixes
        from app.models.email_details import EmailDetails
        automatic_requirements = (
            get_db_session().query(Requirement).join(
                EmailDetails, Requirement.requirement_id == EmailDetails.requirement_id
            )
            .filter(
                Requirement.requirement_id.isnot(None),
                Requirement.is_manual_requirement == False,
                ~EmailDetails.email_subject.like('Re:%'),
                ~EmailDetails.email_subject.like('Fw:%'),
                ~EmailDetails.email_subject.like('Fwd:%'),
                ~EmailDetails.email_subject.like('Forward:%')
            )
            .order_by(Requirement.created_at.desc())
            .all()
        )
        
        # Combine both sets
        requirements = manual_requirements + automatic_requirements
        
        current_app.logger.info(f"Found {len(manual_requirements)} manual requirements and {len(automatic_requirements)} automatic requirements")
        
        # Helper to get latest related EmailDetails fields (subject, email_id)
        def _latest_email_fields(req):
            try:
                email_recs = getattr(req, 'email_details', []) or []
                if not email_recs:
                    return None, None
                latest = max(
                    email_recs,
                    key=lambda e: (e.received_datetime or e.created_at or req.created_at or datetime.min)
                )
                return latest.email_subject, latest.email_id
            except Exception:
                return None, None

        result = []
        for r in requirements:
            email_subject, email_id = _latest_email_fields(r)
            result.append({
                'id': r.requirement_id,
                'request_id': r.request_id,
                'job_title': r.job_title,
                'email_subject': email_subject,
                'company_name': format_enum_for_display(r.company_name),
                'department': format_enum_for_display(r.department),
                'location': r.location,
                'shift': format_enum_for_display(r.shift),
                'job_type': format_enum_for_display(r.job_type),
                'hiring_manager': r.hiring_manager,
                'experience_range': r.experience_range,
                'skills_required': getattr(r, 'skills_required', None),
                'minimum_qualification': r.minimum_qualification,
                'number_of_positions': r.number_of_positions,
                'budget_ctc': r.budget_ctc,
                'priority': format_enum_for_display(r.priority),
                'tentative_doj': r.tentative_doj.isoformat() if r.tentative_doj else None,
                'additional_remarks': r.additional_remarks,
                'assigned_to': get_db_session().query(User).filter_by(user_id=r.user_id).first().username if r.user_id else None,
                'assigned_recruiters': _get_assigned_recruiters_for_requirement(r.requirement_id),
                'is_manual_requirement': r.is_manual_requirement,
                'created_at': r.created_at.isoformat() if r.created_at else None,
                'updated_at': r.updated_at.isoformat() if r.updated_at else None,
                'email_id': email_id
            })
        
        current_app.logger.info("Successfully converted requirements to JSON")
        return jsonify(result)
    except Exception as e:
        current_app.logger.error(f"Error fetching requirements data: {str(e)}")
        return jsonify({'error': 'Failed to fetch requirements data'}), 500

@api_bp.route('/requirements/<int:id>', methods=['GET'])
def get_requirement(id):
    """Get a specific requirement by ID"""
    try:
        requirement = get_db_session().query(Requirement).filter_by(id=id).first_or_404()
        return jsonify(requirement.to_dict())
    except Exception as e:
        current_app.logger.error(f"Error fetching requirement: {str(e)}")
        return jsonify({'error': 'Failed to fetch requirement'}), 500

def check_duplicate_requirement(data):
    """Check if a requirement with similar details already exists"""
    try:
        # Check for exact matches on key fields - fields are now strings, no conversion needed
        exact_matches = get_db_session().query(Requirement).filter(
            Requirement.job_title == data.get('job_title'),
            Requirement.company_name == data.get('company_name'),
            Requirement.department == data.get('department'),
            Requirement.location == data.get('location')
        ).all()
        
        if exact_matches:
            return {
                'is_duplicate': True,
                'duplicates': [req.to_dict() for req in exact_matches],
                'match_type': 'exact'
            }
        
        # Check for similar matches (fuzzy matching)
        similar_matches = []
        
        # Check by job title and company (most important fields)
        title_company_matches = get_db_session().query(Requirement).filter(
            Requirement.job_title == data.get('job_title'),
            Requirement.company_name == data.get('company_name')
        ).all()
        
        if title_company_matches:
            similar_matches.extend(title_company_matches)
        
        # Check by job title and department
        title_dept_matches = get_db_session().query(Requirement).filter(
            Requirement.job_title == data.get('job_title'),
            Requirement.department == data.get('department')
        ).all()
        
        if title_dept_matches:
            similar_matches.extend(title_dept_matches)
        
        # Remove duplicates from similar_matches
        seen_ids = set()
        unique_similar_matches = []
        for req in similar_matches:
            if req.requirement_id not in seen_ids:
                seen_ids.add(req.requirement_id)
                unique_similar_matches.append(req)
        
        if unique_similar_matches:
            return {
                'is_duplicate': True,
                'duplicates': [req.to_dict() for req in unique_similar_matches],
                'match_type': 'similar'
            }
        
        return {
            'is_duplicate': False,
            'duplicates': [],
            'match_type': None
        }
        
    except Exception as e:
        current_app.logger.error(f"Error checking for duplicates: {str(e)}")
        return {
            'is_duplicate': False,
            'duplicates': [],
            'match_type': None
        }

@api_bp.route('/requirements/check-duplicate', methods=['POST'])
def check_requirement_duplicate():
    """Check if a requirement is a duplicate before creating"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        duplicate_check = check_duplicate_requirement(data)
        
        return jsonify({
            'success': True,
            'duplicate_check': duplicate_check
        }), 200
        
    except Exception as e:
        current_app.logger.error(f"Error checking requirement duplicate: {str(e)}")
        return jsonify({'error': 'Failed to check for duplicates'}), 500

@api_bp.route('/requirements', methods=['POST'])
def create_requirement():
    """Create a new requirement manually"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        # Check for duplicates first
        duplicate_check = check_duplicate_requirement(data)

        # If duplicates found, return them for user confirmation
        if duplicate_check['is_duplicate']:
            return jsonify({
                'success': False,
                'has_duplicates': True,
                'duplicate_check': duplicate_check,
                'message': f"Found {len(duplicate_check['duplicates'])} similar requirement(s). Please review before proceeding."
            }), 200

        # Create new requirement - fields are now strings, no conversion needed
        new_requirement = Requirement(
            job_title=data.get('job_title'),
            department=data.get('department'),
            location=data.get('location'),
            shift=data.get('shift'),
            job_type=data.get('job_type'),
            hiring_manager=data.get('hiring_manager'),
            experience_range=data.get('experience_range'),
            minimum_qualification=data.get('minimum_qualification'),
            number_of_positions=data.get('number_of_positions'),
            budget_ctc=data.get('budget_ctc'),
            priority=data.get('priority'),
            tentative_doj=datetime.strptime(data.get('tentative_doj'), '%Y-%m-%d').date() if data.get('tentative_doj') else None,
            additional_remarks=data.get('additional_remarks'),
            company_name=data.get('company_name'),
            notes=data.get('notes'),
            is_manual_requirement=True,  # Mark as manual requirement
            # JD fields
            job_description=data.get('job_description'),
            jd_path=data.get('jd_path'),
            job_file_name=data.get('job_file_name')
        )
        
        get_db_session().add(new_requirement)
        get_db_session().commit()
        
        current_app.logger.info(f"Created new manual requirement: {new_requirement.requirement_id} with is_manual_requirement={new_requirement.is_manual_requirement}")
        return jsonify({
            'success': True,
            'message': 'Requirement created successfully',
            'requirement': new_requirement.to_dict()
        }), 201
        
    except Exception as e:
        get_db_session().rollback()
        current_app.logger.error(f"Error creating requirement: {str(e)}")
        return jsonify({'error': 'Failed to create requirement'}), 500

@api_bp.route('/requirements/force-create', methods=['POST'])
def force_create_requirement():
    """Force create a new requirement even if duplicates exist"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        # Create new requirement - fields are now strings, no conversion needed
        new_requirement = Requirement(
            job_title=data.get('job_title'),
            department=data.get('department'),
            location=data.get('location'),
            shift=data.get('shift'),
            job_type=data.get('job_type'),
            hiring_manager=data.get('hiring_manager'),
            experience_range=data.get('experience_range'),
            minimum_qualification=data.get('minimum_qualification'),
            number_of_positions=data.get('number_of_positions'),
            budget_ctc=data.get('budget_ctc'),
            priority=data.get('priority'),
            tentative_doj=datetime.strptime(data.get('tentative_doj'), '%Y-%m-%d').date() if data.get('tentative_doj') else None,
            additional_remarks=data.get('additional_remarks'),
            company_name=data.get('company_name'),
            notes=data.get('notes'),
            is_manual_requirement=True,  # Mark as manual requirement
            # JD fields
            job_description=data.get('job_description'),
            jd_path=data.get('jd_path'),
            job_file_name=data.get('job_file_name')
        )
        
        get_db_session().add(new_requirement)
        get_db_session().commit()
        
        current_app.logger.info(f"Force created new manual requirement: {new_requirement.requirement_id} with is_manual_requirement={new_requirement.is_manual_requirement}")
        return jsonify({
            'success': True,
            'message': 'Requirement created successfully (duplicates ignored)',
            'requirement': new_requirement.to_dict()
        }), 201
        
    except Exception as e:
        get_db_session().rollback()
        current_app.logger.error(f"Error force creating requirement: {str(e)}")
        return jsonify({'error': 'Failed to create requirement'}), 500

@api_bp.route('/requirements/bulk-upload', methods=['POST'])
def bulk_upload_requirements():
    """Upload multiple requirements from Excel or CSV file"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']



        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        # Check file extension
        allowed_extensions = {'xlsx', 'xls', 'csv'}
        file_extension = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''

        if file_extension not in allowed_extensions:
            return jsonify({'error': f'File type not allowed. Allowed types: {", ".join(allowed_extensions)}'}), 400

        # Check file size (10MB limit)
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning

        if file_size > 10 * 1024 * 1024:  # 10MB
            return jsonify({'error': 'File size too large. Maximum size is 10MB'}), 400

        # Read file based on extension
        import pandas as pd
        from io import BytesIO, StringIO

        try:
            if file_extension == 'csv':
                # Read CSV file
                content = file.read().decode('utf-8')
                df = pd.read_csv(StringIO(content))
            else:
                # Read Excel file
                df = pd.read_excel(file, engine='openpyxl')
        except Exception as e:
            return jsonify({'error': f'Error reading file: {str(e)}'}), 400
        
        # Validate required columns
        required_columns = ['Job Title', 'Company', 'Department', 'Location']
        missing_columns = [col for col in required_columns if col not in df.columns]
        
        if missing_columns:
            return jsonify({
                'error': f'Missing required columns: {", ".join(missing_columns)}. Required columns: {", ".join(required_columns)}'
            }), 400
        
        # Optional columns mapping
        optional_columns = {
            'Shift': 'shift',
            'Job Type': 'job_type',
            'Hiring Manager': 'hiring_manager',
            'Experience Range': 'experience_range',
            'Minimum Qualification': 'minimum_qualification',
            'Number of Positions': 'number_of_positions',
            'Budget CTC': 'budget_ctc',
            'Priority': 'priority',
            'Tentative DOJ': 'tentative_doj',
            'Additional Remarks': 'additional_remarks',
            'Assigned To': 'assigned_to'
        }
        
        # Get JD info if provided
        jd_info = None
        if 'jd_info' in request.form:
            try:
                jd_info = json.loads(request.form['jd_info'])
            except json.JSONDecodeError:
                current_app.logger.warning("Invalid JD info JSON provided")
        
        # Pre-generate unique request_ids for all rows to avoid race conditions
        def generate_unique_request_ids(count):
            """Generate a list of unique request_ids atomically"""
            request_ids = []

            # Get the current max request_id (string like "Req003") with exclusive lock to prevent race conditions
            latest_req = get_db_session().query(Requirement).with_for_update().order_by(Requirement.request_id.desc()).first()

            # Derive starting number from the string request_id, not the UUID primary key
            if latest_req and getattr(latest_req, 'request_id', None):
                try:
                    start_num = int(str(latest_req.request_id).replace('Req', '')) + 1
                except (ValueError, AttributeError):
                    start_num = 1
            else:
                start_num = 1

            print(f"Starting request_id generation from: Req{start_num:03d}")  # Debug log

            # Generate sequential IDs starting from the next available number
            current_num = start_num
            for _ in range(count):
                # Double-check if the generated request_id already exists (extra safety)
                while True:
                    candidate_id = f"Req{current_num:03d}"
                    existing = get_db_session().query(Requirement).filter_by(request_id=candidate_id).first()
                    if not existing:
                        request_ids.append(candidate_id)
                        current_num += 1
                        break
                    else:
                        print(f"Request ID {candidate_id} already exists, skipping...")  # Debug log
                        current_num += 1

            return request_ids

        # Filter out completely empty rows first
        df = df.dropna(how='all')

        # Generate all request_ids upfront (only for non-empty rows)
        total_rows = len(df)
        print(f"Total rows after filtering empty rows: {total_rows}")  # Debug log
        request_ids = generate_unique_request_ids(total_rows)
        print(f"Generated request_ids: {request_ids}")  # Debug log

        # Validate that we have enough request_ids
        if len(request_ids) != total_rows:
            print(f"Warning: Generated {len(request_ids)} request_ids but need {total_rows}")
            return jsonify({'error': f'Failed to generate sufficient request_ids: got {len(request_ids)}, needed {total_rows}'}), 500

        # Process each row
        results = []
        success_count = 0
        error_count = 0
        request_id_index = 0

        for index, row in df.iterrows():
            print(f"Processing row {index + 1}, request_id_index: {request_id_index}, total_request_ids: {len(request_ids)}")  # Debug log
            try:
                # Validate required fields
                if pd.isna(row['Job Title']) or str(row['Job Title']).strip() == '':
                    results.append({
                        'row': index + 1,
                        'status': 'error',
                        'message': 'Job Title is required'
                    })
                    error_count += 1
                    request_id_index += 1
                    continue
                
                if pd.isna(row['Company']) or str(row['Company']).strip() == '':
                    results.append({
                        'row': index + 1,
                        'status': 'error',
                        'message': 'Company is required'
                    })
                    error_count += 1
                    request_id_index += 1
                    continue

                if pd.isna(row['Department']) or str(row['Department']).strip() == '':
                    results.append({
                        'row': index + 1,
                        'status': 'error',
                        'message': 'Department is required'
                    })
                    error_count += 1
                    request_id_index += 1
                    continue
                
                if pd.isna(row['Location']) or str(row['Location']).strip() == '':
                    results.append({
                        'row': index + 1,
                        'status': 'error',
                        'message': 'Location is required'
                    })
                    error_count += 1
                    request_id_index += 1
                    continue

                # Use pre-generated unique request_id with bounds checking
                if request_id_index >= len(request_ids):
                    results.append({
                        'row': index + 1,
                        'status': 'error',
                        'message': f'Index out of range: {request_id_index} >= {len(request_ids)}'
                    })
                    error_count += 1
                    continue

                request_id = request_ids[request_id_index]
                request_id_index += 1
                print(f"Assigned request_id: {request_id}")  # Debug log

                # Prepare data for requirement creation
                requirement_data = {
                    'request_id': request_id,
                    'job_title': str(row['Job Title']).strip(),
                    'company_name': str(row['Company']).strip(),
                    'department': str(row['Department']).strip(),
                    'location': str(row['Location']).strip(),
                    'is_manual_requirement': True,
                    # JD fields (include if provided)
                    'job_description': jd_info.get('job_description') if jd_info else None,
                    'jd_path': jd_info.get('jd_path') if jd_info else None,
                    'job_file_name': jd_info.get('job_file_name') if jd_info else None
                }
                
                # Add optional fields
                for excel_col, db_field in optional_columns.items():
                    if excel_col in df.columns and not pd.isna(row[excel_col]) and str(row[excel_col]).strip() != '':
                        value = str(row[excel_col]).strip()
                        
                        # Handle special cases
                        if excel_col == 'Number of Positions':
                            try:
                                value = int(value)
                            except ValueError:
                                value = None
                        elif excel_col == 'Tentative DOJ':
                            try:
                                value = datetime.strptime(value, '%Y-%m-%d').date()
                            except ValueError:
                                try:
                                    value = pd.to_datetime(value).date()
                                except:
                                    value = None
                        
                        if value is not None:
                            requirement_data[db_field] = value

                # Fields are now strings, no enum conversion needed
                # Just clean up the department and company_name values
                if 'department' in requirement_data and isinstance(requirement_data['department'], str):
                    # Allow plain strings by mapping spaces to underscores
                    requirement_data['department'] = requirement_data['department'].replace(' ', '_')
                if 'company_name' in requirement_data and isinstance(requirement_data['company_name'], str):
                    requirement_data['company_name'] = requirement_data['company_name'].replace(' ', '_')
                
                # Check for duplicates before creating
                duplicate_check = check_duplicate_requirement(requirement_data)
                
                if duplicate_check['is_duplicate']:
                    results.append({
                        'row': index + 1,
                        'status': 'duplicate',
                        'job_title': requirement_data['job_title'],
                        'duplicates': duplicate_check['duplicates'],
                        'match_type': duplicate_check['match_type'],
                        'message': f"Found {len(duplicate_check['duplicates'])} similar requirement(s)"
                    })
                    error_count += 1
                    request_id_index += 1
                else:
                    # Create requirement
                    new_requirement = Requirement(**requirement_data)
                    get_db_session().add(new_requirement)
                    get_db_session().commit()
                    
                    results.append({
                        'row': index + 1,
                        'status': 'success',
                        'request_id': request_id,
                        'job_title': requirement_data['job_title']
                    })
                    success_count += 1
                
            except Exception as e:
                get_db_session().rollback()
                current_app.logger.error(f"Error processing row {index + 1}: {str(e)}")
                results.append({
                    'row': index + 1,
                    'status': 'error',
                    'message': f'Error creating requirement: {str(e)}'
                })
                error_count += 1
        
        return jsonify({
            'success': True,
            'message': f'Bulk upload completed. {success_count} requirements created, {error_count} errors.',
            'results': results,
            'summary': {
                'total_rows': len(df),
                'success_count': success_count,
                'error_count': error_count
            }
        }), 200
        
    except Exception as e:
        current_app.logger.error(f"Error in bulk upload: {str(e)}")
        return jsonify({'error': f'Failed to process bulk upload: {str(e)}'}), 500


@api_bp.route('/requirements/upload-jd', methods=['POST'])
def upload_job_description():
    """Upload job description file (PDF or DOCX)"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400

        file = request.files['file']

        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        # Check file extension
        allowed_extensions = {'pdf', 'docx'}
        file_extension = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''

        if file_extension not in allowed_extensions:
            return jsonify({'error': f'File type not allowed. Only PDF and DOCX files are allowed.'}), 400

        # Check file size (10MB limit)
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning

        if file_size > 10 * 1024 * 1024:  # 10MB
            return jsonify({'error': 'File size too large. Maximum size is 10MB'}), 400

        # Generate unique filename
        import uuid
        import os
        from datetime import datetime
        from werkzeug.utils import secure_filename

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        unique_id = str(uuid.uuid4())[:8]
        original_filename = secure_filename(file.filename)
        safe_filename = f"JD_{timestamp}_{unique_id}.{file_extension}"

        # Create uploads directory if it doesn't exist
        upload_dir = os.path.join(current_app.root_path, 'uploads', 'job_descriptions')
        os.makedirs(upload_dir, exist_ok=True)

        # Save file
        file_path = os.path.join(upload_dir, safe_filename)
        file.save(file_path)

        # Extract text from the file (optional - for search purposes)
        job_description_text = None
        try:
            if file_extension == 'pdf':
                # Extract text from PDF
                import PyPDF2
                with open(file_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    text_content = []
                    for page in pdf_reader.pages:
                        text_content.append(page.extract_text())
                    job_description_text = '\n'.join(text_content)
            elif file_extension == 'docx':
                # Extract text from DOCX
                from docx import Document
                doc = Document(file_path)
                text_content = []
                for paragraph in doc.paragraphs:
                    text_content.append(paragraph.text)
                job_description_text = '\n'.join(text_content)
        except Exception as e:
            current_app.logger.warning(f"Could not extract text from JD file: {str(e)}")
            # Continue without text extraction

        return jsonify({
            'success': True,
            'message': 'Job description uploaded successfully',
            'file_info': {
                'original_filename': original_filename,
                'stored_filename': safe_filename,
                'file_path': file_path,
                'file_size': file_size,
                'file_extension': file_extension,
                'job_description_text': job_description_text
            }
        }), 200

    except Exception as e:
        current_app.logger.error(f"Error uploading job description: {str(e)}")
        return jsonify({'error': f'Failed to upload job description: {str(e)}'}), 500


@api_bp.route('/requirements/update-jd', methods=['POST'])
def update_requirement_jd():
    """Update job description for an existing requirement"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'}), 400

        file = request.files['file']
        requirement_id = request.form.get('requirement_id')

        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'}), 400

        if not requirement_id:
            return jsonify({'success': False, 'error': 'No requirement_id provided'}), 400

        # Check file extension
        allowed_extensions = {'pdf', 'docx'}
        file_extension = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''

        if file_extension not in allowed_extensions:
            return jsonify({'success': False, 'error': f'File type not allowed. Only PDF and DOCX files are allowed.'}), 400

        # Check file size (10MB limit)
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning

        if file_size > 10 * 1024 * 1024:  # 10MB
            return jsonify({'success': False, 'error': 'File size too large. Maximum size is 10MB'}), 400

        # Find the requirement - handle both request_id and requirement_id
        if requirement_id.startswith('Req'):
            # It's a request_id, so query by request_id
            requirement = get_db_session().query(Requirement).filter_by(request_id=requirement_id).first()
        else:
            # It's a UUID requirement_id
            requirement = get_db_session().query(Requirement).filter_by(requirement_id=requirement_id).first()
        
        if not requirement:
            return jsonify({'success': False, 'error': 'Requirement not found'}), 404

        # Generate unique filename
        import uuid
        import os
        from datetime import datetime
        from werkzeug.utils import secure_filename

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        unique_id = str(uuid.uuid4())[:8]
        original_filename = secure_filename(file.filename)
        safe_filename = f"JD_{timestamp}_{unique_id}.{file_extension}"

        # Create uploads directory if it doesn't exist
        upload_dir = os.path.join(current_app.root_path, 'uploads', 'job_descriptions')
        os.makedirs(upload_dir, exist_ok=True)

        # Save file
        file_path = os.path.join(upload_dir, safe_filename)
        file.save(file_path)

        # Extract text from the file (optional - for search purposes)
        job_description_text = None
        try:
            if file_extension == 'pdf':
                # Extract text from PDF
                import PyPDF2
                with open(file_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    text_content = []
                    for page in pdf_reader.pages:
                        text_content.append(page.extract_text())
                    job_description_text = '\n'.join(text_content)
            elif file_extension == 'docx':
                # Extract text from DOCX
                from docx import Document
                doc = Document(file_path)
                text_content = []
                for paragraph in doc.paragraphs:
                    text_content.append(paragraph.text)
                job_description_text = '\n'.join(text_content)
        except Exception as e:
            current_app.logger.warning(f"Could not extract text from JD file: {str(e)}")
            # Continue without text extraction

        # Update the requirement with JD information
        requirement.job_description = job_description_text
        requirement.jd_path = file_path  # Store the full path, not just filename
        requirement.job_file_name = original_filename
        requirement.updated_at = datetime.utcnow()

        get_db_session().commit()

        current_app.logger.info(f"Updated JD for requirement {requirement.request_id} (ID: {requirement.requirement_id}): {original_filename}")

        return jsonify({
            'success': True,
            'message': 'Job description updated successfully',
            'file_info': {
                'original_filename': original_filename,
                'stored_filename': safe_filename,
                'file_path': file_path,
                'file_size': file_size,
                'file_extension': file_extension,
                'job_description_text': job_description_text
            }
        }), 200

    except Exception as e:
        get_db_session().rollback()
        current_app.logger.error(f"Error updating job description: {str(e)}")
        return jsonify({'success': False, 'error': f'Failed to update job description: {str(e)}'}), 500


@api_bp.route('/upload-profiles', methods=['POST'])
def upload_profiles():
    """Upload candidate profiles from file"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'}), 400
        
        file = request.files['file']
        request_id = request.form.get('request_id')
        
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'}), 400
        
        if not request_id:
            return jsonify({'success': False, 'error': 'No request_id provided'}), 400
        
        # Check file extension
        allowed_extensions = {'xlsx', 'xls', 'csv'}
        file_extension = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''
        
        if file_extension not in allowed_extensions:
            return jsonify({'success': False, 'error': f'File type not allowed. Allowed types: {", ".join(allowed_extensions)}'}), 400
        
        # Check file size (10MB limit)
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning
        
        if file_size > 10 * 1024 * 1024:  # 10MB
            return jsonify({'success': False, 'error': 'File size too large. Maximum size is 10MB'}), 400
        
        # Save file temporarily
        import tempfile
        import os
        from werkzeug.utils import secure_filename
        
        filename = secure_filename(file.filename)
        temp_dir = os.path.join(current_app.root_path, 'uploads', 'temp')
        os.makedirs(temp_dir, exist_ok=True)
        
        temp_path = os.path.join(temp_dir, f"{request_id}_{filename}")
        file.save(temp_path)
        
        try:
            # Process the file using EmailProcessor
            email_processor = EmailProcessor()
            
            # Read file content based on type
            if file_extension in ['xlsx', 'xls']:
                import pandas as pd
                df = pd.read_excel(temp_path)
            else:  # csv
                import pandas as pd
                df = pd.read_csv(temp_path)
            
            # Convert DataFrame to HTML for processing
            html_content = df.to_html(index=False)
            
            # Extract profiles using existing logic
            profiles = email_processor.extract_profiles_from_html(html_content)
            
            if not profiles:
                return jsonify({'success': False, 'error': 'No valid profiles found in the uploaded file'}), 400
            
            # Create profiles in database
            created_count = 0
            for profile_data in profiles:
                try:
                    # Add email_id and request_id to profile data
                    profile_data['email_id'] = f"upload_{request_id}_{created_count}"
                    
                    # Create profile using existing logic
                    profile = email_processor._create_or_update_profile(profile_data, f"upload_{request_id}")
                    if profile:
                        created_count += 1
                except Exception as e:
                    current_app.logger.error(f"Error creating profile: {str(e)}")
                    continue
            
            # Clean up temporary file
            os.remove(temp_path)
            
            return jsonify({
                'success': True,
                'message': f'Successfully uploaded {created_count} profiles',
                'profiles_created': created_count,
                'request_id': request_id
            })
            
        except Exception as e:
            # Clean up temporary file on error
            if os.path.exists(temp_path):
                os.remove(temp_path)
            raise e
        
    except Exception as e:
        current_app.logger.error(f"Error uploading profiles: {str(e)}")
        return jsonify({'success': False, 'error': f'Failed to upload profiles: {str(e)}'}), 500

# Unified Authentication Endpoints
@api_bp.route('/login', methods=['POST'])
def login():
    """Unified login for both admin and recruiter users with domain isolation"""
    try:
        current_app.logger.info("Login endpoint called")
        data = request.get_json()
        current_app.logger.info(f"Received data: {data}")
        
        username = data.get('username')
        password = data.get('password')
        
        current_app.logger.info(f"Username: {username}, Password: {password}")
        
        if not username or not password:
            current_app.logger.warning("Missing username or password")
            return jsonify({
                'status': 'error',
                'message': 'Username and password are required'
            }), 400
        
        # Get domain from custom header or fallback to detection
        domain = request.headers.get('X-Original-Domain')
        if not domain:
            domain = request.headers.get('X-Domain')
        
        # Ensure domain database isolation before authentication
        if not ensure_domain_isolation():
            current_app.logger.error(f"Failed to ensure domain database isolation for domain: {domain}")
            return jsonify({
                'status': 'error',
                'message': 'Database not available for this domain'
            }), 503
        
        # Get domain-specific database session
        from flask import g
        if not hasattr(g, 'db_session') or g.db_session is None:
            current_app.logger.error("No database session available for domain")
            return jsonify({
                'status': 'error',
                'message': 'Database session not available for this domain'
            }), 503
        
        # Find user by username in domain-specific database
        user = g.db_session.query(User).filter_by(username=username).first()
        current_app.logger.info(f"User found: {user}")
        
        if not user:
            current_app.logger.warning(f"No user found with username: {username}")
            return jsonify({
                'status': 'error',
                'message': 'Invalid username or password'
            }), 401
        
        # Check password using the new hashing method
        if not user.check_password(password):
            current_app.logger.warning(f"Password mismatch for user: {username}")
            return jsonify({
                'status': 'error',
                'message': 'Invalid username or password'
            }), 401
        
        # Create JWT token
        access_token = create_access_token(identity=username)
        
        current_app.logger.info(f"Login successful for user: {username}")
        return jsonify({
            'status': 'success',
            'message': 'Login successful',
            'access_token': access_token,
            'user': user.to_dict()
        })
        
    except Exception as e:
        current_app.logger.error(f"Error in login: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to login'
        }), 500

@api_bp.route('/signup', methods=['POST'])
def signup():
    """Unified signup for both admin and recruiter users with domain isolation"""
    try:
        data = request.get_json()
        username = data.get('username')
        password = data.get('password')
        email = data.get('email')
        full_name = data.get('full_name')
        phone_number = data.get('phone_number')
        role = data.get('role', 'recruiter')  # Default to recruiter if not specified
        
        if not username or not password or not full_name or not email:
            return jsonify({
                'status': 'error',
                'message': 'username, password, full_name, and email are required'
            }), 400
        
        # Validate role
        if role not in ['admin', 'recruiter']:
            return jsonify({
                'status': 'error',
                'message': 'Invalid role. Must be either "admin" or "recruiter"'
            }), 400
        
        # Validate phone number (optional but if present must be 10 digits)
        if phone_number is not None:
            phone_str = str(phone_number)
            if not phone_str.isdigit() or len(phone_str) != 10:
                return jsonify({
                    'status': 'error',
                    'message': 'phone_number must be a 10-digit number'
                }), 400
        
        # Get domain from custom header or fallback to detection
        domain = request.headers.get('X-Original-Domain')
        if not domain:
            domain = request.headers.get('X-Domain')
        
        # Ensure domain database isolation before user creation
        if not ensure_domain_isolation():
            current_app.logger.error(f"Failed to ensure domain database isolation for domain: {domain}")
            return jsonify({
                'status': 'error',
                'message': 'Database not available for this domain'
            }), 503
        
        # Get domain-specific database session
        from flask import g
        if not hasattr(g, 'db_session') or g.db_session is None:
            current_app.logger.error("No database session available for domain")
            return jsonify({
                'status': 'error',
                'message': 'Database session not available for this domain'
            }), 503
        
        # Check if username already exists in domain-specific database
        existing_user = g.db_session.query(User).filter_by(username=username).first()
        if existing_user:
            return jsonify({
                'status': 'error',
                'message': 'Username already exists'
            }), 409
        
        # Check if email already exists in domain-specific database (excluding temporary users)
        existing_email = g.db_session.query(User).filter(User.email == email, User.username.notlike('temp_%')).first()
        if existing_email:
            return jsonify({
                'status': 'error',
                'message': 'Email already exists'
            }), 409
        
        # Check if there's a temporary user with this email in domain-specific database (indicating OTP verification)
        temp_user = g.db_session.query(User).filter(User.email == email, User.username.like('temp_%')).first()
        if not temp_user:
            # Check if there's already a permanent user with this email in domain-specific database
            permanent_user = g.db_session.query(User).filter(User.email == email, User.username.notlike('temp_%')).first()
            if permanent_user:
                return jsonify({
                    'status': 'error',
                    'message': 'An account with this email already exists. Please try logging in instead.'
                }), 409
            else:
                return jsonify({
                    'status': 'error',
                    'message': 'Please verify your email address first by sending and verifying an OTP'
                }), 400
        
        # Update the temporary user with the actual user data
        from app.models.user import UserRoleEnum
        role_value = UserRoleEnum.admin if role == 'admin' else UserRoleEnum.recruiter
        
        # Update the temporary user with the real user data
        temp_user.username = username
        temp_user.full_name = full_name
        temp_user.role = role_value
        temp_user.otp = None  # Clear OTP
        temp_user.otp_expiry_time = None  # Clear OTP expiry
        
        if phone_number is not None:
            # store as numeric field; SQLAlchemy Numeric accepts str or int
            temp_user.phone_number = int(phone_number)
        
        # Set the new password
        temp_user.set_password(password)
        
        # Commit the changes using domain-specific database session
        g.db_session.commit()
        
        return jsonify({
            'status': 'success',
            'message': f'{role.capitalize()} account created successfully',
            'user': temp_user.to_dict()
        }), 201
        
    except Exception as e:
        current_app.logger.error(f"Error in signup: {str(e)}")
        if hasattr(g, 'db_session') and g.db_session is not None:
            g.db_session.rollback()
        
        # Check if it's a unique constraint violation
        if 'UniqueViolation' in str(e) or 'duplicate key' in str(e):
            if 'email' in str(e):
                return jsonify({
                    'status': 'error',
                    'message': 'An account with this email already exists. Please try logging in instead.'
                }), 409
            elif 'username' in str(e):
                return jsonify({
                    'status': 'error',
                    'message': 'Username already exists. Please choose a different username.'
                }), 409
        
        return jsonify({
            'status': 'error',
            'message': 'Failed to create account'
        }), 500

@api_bp.route('/auth/current-user', methods=['GET'])
@require_jwt_domain_auth
def get_current_user():
    """Get current authenticated user information with JWT and domain isolation"""
    try:
        # User is already authenticated and available in request context
        user = getattr(request, 'current_user', None)
        if user:
            return jsonify({
                'status': 'success',
                'user': user.to_dict()
            })
        else:
            return jsonify({
                'status': 'error',
                'message': 'User not found'
            }), 404
            
    except Exception as e:
        current_app.logger.error(f"Error getting current user: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to get current user'
        }), 500

@api_bp.route('/auth/send-otp', methods=['POST'])
def send_otp():
    """Send OTP to email for verification"""
    try:
        data = request.get_json()
        email = data.get('email')
        
        if not email:
            return jsonify({
                'status': 'error',
                'message': 'Email is required'
            }), 400
        
        # Check if email already exists
        existing_user = get_db_session().query(User).filter_by(email=email).first()
        if existing_user:
            return jsonify({
                'status': 'error',
                'message': 'Email already exists'
            }), 409
        
        # Generate 6-digit OTP
        import random
        otp = random.randint(100000, 999999)
        
        # Store OTP in database (we'll create a temporary record or use a separate OTP table)
        # For now, we'll store it in a temporary way
        from datetime import datetime, timedelta
        otp_expiry = datetime.utcnow() + timedelta(minutes=10)  # OTP expires in 10 minutes
        
        # For simplicity, we'll create a temporary user record with OTP
        # In production, you might want to use a separate OTP table
        temp_user = User(
            username=f"temp_{email}_{datetime.utcnow().timestamp()}",
            full_name="Temporary User",
            email=email,
            password="temp_password",  # This will be replaced when user completes signup
            otp=otp,
            otp_expiry_time=otp_expiry
        )
        temp_user.set_password("temp_password")
        
        get_db_session().add(temp_user)
        get_db_session().commit()
        
        # Send OTP via email
        try:
            from app.services.email_service import EmailService
            email_service = EmailService()
            
            subject = "Email Verification OTP"
            html_content = f"""
            <html>
            <body>
                <h2>Email Verification</h2>
                <p>Your verification code is: <strong>{otp}</strong></p>
                <p>This code will expire in 10 minutes.</p>
                <p>If you did not request this code, please ignore this email.</p>
            </body>
            </html>
            """
            
            email_sent = email_service.send_email(email, subject, html_content)
            
            if email_sent:
                return jsonify({
                    'status': 'success',
                    'message': 'OTP sent successfully'
                }), 200
            else:
                # If email sending fails, clean up the temp user
                get_db_session().delete(temp_user)
                get_db_session().commit()
                return jsonify({
                    'status': 'error',
                    'message': 'Failed to send OTP email'
                }), 500
                
        except Exception as e:
            # If email sending fails, clean up the temp user
            get_db_session().delete(temp_user)
            get_db_session().commit()
            current_app.logger.error(f"Error sending OTP email: {str(e)}")
            return jsonify({
                'status': 'error',
                'message': 'Failed to send OTP email'
            }), 500
        
    except Exception as e:
        current_app.logger.error(f"Error in send_otp: {str(e)}")
        get_db_session().rollback()
        return jsonify({
            'status': 'error',
            'message': 'Failed to send OTP'
        }), 500

@api_bp.route('/auth/verify-otp', methods=['POST'])
def verify_otp():
    """Verify OTP for email verification"""
    try:
        data = request.get_json()
        email = data.get('email')
        otp = data.get('otp')
        
        if not email or not otp:
            return jsonify({
                'status': 'error',
                'message': 'Email and OTP are required'
            }), 400
        
        # Find the temporary user with the email and OTP
        temp_user = get_db_session().query(User).filter_by(email=email, otp=otp).first()
        
        if not temp_user:
            return jsonify({
                'status': 'error',
                'message': 'Invalid OTP'
            }), 400
        
        # Check if OTP has expired
        from datetime import datetime
        if temp_user.otp_expiry_time and temp_user.otp_expiry_time < datetime.utcnow():
            get_db_session().delete(temp_user)
            get_db_session().commit()
            return jsonify({
                'status': 'error',
                'message': 'OTP has expired'
            }), 400
        
        # OTP is valid, we can mark it as verified
        # In a real implementation, you might want to store this verification status
        # For now, we'll just return success
        return jsonify({
            'status': 'success',
            'message': 'OTP verified successfully'
        }), 200
        
    except Exception as e:
        current_app.logger.error(f"Error in verify_otp: {str(e)}")
        get_db_session().rollback()
        return jsonify({
            'status': 'error',
            'message': 'Failed to verify OTP'
        }), 500

@api_bp.route('/users', methods=['GET'])
@require_domain_auth
def get_users():
    """Get all users (admin only)"""
    try:
        # Get all users from User table
        users = get_db_session().query(User).all()
        
        users_data = []
        for user in users:
            user_dict = {
                'id': str(user.user_id),
                'username': user.username,
                'full_name': user.full_name,
                'email': user.email,
                'phone_number': str(user.phone_number) if user.phone_number else None,
                'role': user.role.value if user.role else None,
                'created_at': user.created_at.isoformat() if user.created_at else None,
                'updated_at': user.updated_at.isoformat() if user.updated_at else None
            }
            users_data.append(user_dict)
        
        return jsonify(users_data)
        
    except Exception as e:
        current_app.logger.error(f"Error fetching users: {str(e)}")
        return jsonify({'error': 'Failed to fetch users'}), 500

@api_bp.route('/users/<int:user_id>', methods=['DELETE'])
def delete_user(user_id):
    """Delete a user (admin only)"""
    try:
        user = get_db_session().query(User).filter_by(id=str(user_id)).first()
        if not user:
            return jsonify({'success': False, 'message': 'User not found'}), 404
        
        # Don't allow deleting the last admin
        if user.role == 'admin':
            admin_count = get_db_session().query(User).filter_by(role='admin').count()
            if admin_count <= 1:
                return jsonify({'success': False, 'message': 'Cannot delete the last admin user'}), 400
        
        get_db_session().delete(user)
        get_db_session().commit()
        
        current_app.logger.info(f"Deleted user: {user.username}")
        return jsonify({'success': True, 'message': f'User {user.username} deleted successfully'})
        
    except Exception as e:
        get_db_session().rollback()
        current_app.logger.error(f"Error deleting user: {str(e)}")
        return jsonify({'success': False, 'message': f'Failed to delete user: {str(e)}'}), 500

@api_bp.route('/export-profiles/<string:request_id>', methods=['GET'])
def export_profiles(request_id):
    """Export profiles for a specific request_id to Excel"""
    try:
        # Get profiles for this request_id using the new schema
        from app.models.profile import Profile
        from app.models.requirement import Requirement
        
        # Get the requirement first
        requirement = get_db_session().query(Requirement).filter_by(request_id=request_id).first()
        if not requirement:
            return jsonify({'error': 'Request not found'}), 404
        
        # Get all profiles linked to this requirement
        profiles = get_db_session().query(Profile).filter(
            Profile.requirement_id == requirement.requirement_id,
            Profile.deleted_at.is_(None)
        ).all()
        
        if not profiles:
            return jsonify({'error': 'No profiles found for this request'}), 404
        
        # Convert profiles to dictionaries
        profiles_data = [profile.to_dict() for profile in profiles]
        
        # Get hiring manager name from requirement
        hiring_manager_name = requirement.hiring_manager if requirement else "Hiring Manager"
        
        # Export to formatted document
        from app.services.export_handler import ExportHandler
        export_handler = ExportHandler()
        file_path = export_handler.export_profiles(profiles_data, request_id, hiring_manager_name)
        
        # Get filename from path
        filename = os.path.basename(file_path)
        
        # Debug logging
        current_app.logger.info(f"Exported profiles file created: {file_path}")
        current_app.logger.info(f"File exists: {os.path.exists(file_path)}")
        current_app.logger.info(f"File size: {os.path.getsize(file_path) if os.path.exists(file_path) else 'N/A'}")
        
        return jsonify({
            'success': True,
            'message': f'Successfully exported {len(profiles)} profiles',
            'filename': filename,
            'download_url': f'/exports/{filename}'
        })
        
    except Exception as e:
        current_app.logger.error(f"Error exporting profiles for request {request_id}: {str(e)}")
        return jsonify({'error': f'Failed to export profiles: {str(e)}'}), 500

@api_bp.route('/send-profiles-email/<string:request_id>', methods=['POST'])
def send_profiles_email(request_id):
    """Send profiles email for a specific request_id"""
    try:
        # Handle both JSON and multipart form data
        if request.content_type and 'multipart/form-data' in request.content_type:
            # Handle multipart form data (with file attachments)
            recipient_email = request.form.get('recipient_email')
            recipient_name = request.form.get('recipient_name', 'Hiring Manager')
            cc_email = request.form.get('cc_email', '')  # Optional Cc field
            custom_subject = request.form.get('subject')
            selected_columns = request.form.get('selected_columns', '[]')  # JSON string
            selected_profiles = request.form.get('selected_profiles', '[]')  # JSON string
            
            # Parse JSON strings
            import json
            try:
                selected_columns = json.loads(selected_columns) if selected_columns else []
                selected_profiles = json.loads(selected_profiles) if selected_profiles else []
            except json.JSONDecodeError:
                return jsonify({'error': 'Invalid JSON format in form data'}), 400
            
            # Get file attachments
            attachments = request.files.getlist('attachments')
        else:
            # Handle JSON data (backward compatibility)
            data = request.get_json()
            recipient_email = data.get('recipient_email')
            recipient_name = data.get('recipient_name', 'Hiring Manager')
            cc_email = data.get('cc_email', '')  # Optional Cc field
            custom_subject = data.get('subject')
            selected_columns = data.get('selected_columns', [])  # Selected columns for email template
            selected_profiles = data.get('selected_profiles', [])  # Selected profile student_ids
            attachments = []
        
        if not recipient_email:
            return jsonify({'error': 'Recipient email is required'}), 400
        
        # Get profiles for this request_id using the new schema
        from app.models.profile import Profile
        from app.models.requirement import Requirement
        
        # Get the requirement first
        requirement = get_db_session().query(Requirement).filter_by(request_id=request_id).first()
        if not requirement:
            return jsonify({'error': 'Request not found'}), 404
        
        # Get all profiles linked to this requirement
        all_profiles = get_db_session().query(Profile).filter(
            Profile.requirement_id == requirement.requirement_id,
            Profile.deleted_at.is_(None)
        ).all()
        
        if not all_profiles:
            return jsonify({'error': 'No profiles found for this request'}), 404
        
        # Filter to only selected profiles if provided, otherwise use all
        if selected_profiles:
            profiles = [profile for profile in all_profiles if profile.student_id in selected_profiles]
            current_app.logger.info(f"Filtering profiles: {len(selected_profiles)} selected out of {len(all_profiles)} total")
        else:
            profiles = all_profiles
            current_app.logger.info(f"Using all profiles: {len(profiles)} total")
        
        if not profiles:
            if selected_profiles:
                return jsonify({'error': 'No selected profiles found for this request'}), 404
            else:
                return jsonify({'error': 'No profiles found for this request'}), 404
        
        # Convert profiles to dictionaries
        profiles_data = [profile.to_dict() for profile in profiles]
        
        # Generate email content using the same template as export
        from app.services.export_handler import ExportHandler
        export_handler = ExportHandler()
        email_content = export_handler.generate_email_content(profiles_data, request_id, requirement, recipient_name, selected_columns)
        
        # Send email using email service
        from app.services.email_service import EmailService
        email_service = EmailService()
        
        # Use custom subject if provided, otherwise use default
        subject = custom_subject if custom_subject else f"Candidate Profiles for {requirement.job_title} - {request_id}"
        
        # Send the email with profiles data for Word document generation
        email_sent = email_service.send_email(
            to_email=recipient_email,
            subject=subject,
            html_content=email_content,
            profiles=profiles,
            request_id=request_id,
            requirement=requirement,
            recipient_name=recipient_name,
            cc_email=cc_email,
            attachments=attachments
        )
        
        if email_sent:
            cc_info = f" and Cc'd to {cc_email}" if cc_email else ""
            return jsonify({
                'success': True,
                'message': f'Successfully sent email with {len(profiles)} profiles to {recipient_email}{cc_info}'
            })
        else:
            return jsonify({'error': 'Failed to send email'}), 500
        
    except Exception as e:
        current_app.logger.error(f"Error sending profiles email for request {request_id}: {str(e)}")
        return jsonify({'error': f'Failed to send email: {str(e)}'}), 500
@api_bp.route('/profiles/<string:student_id>', methods=['PUT'])
def update_profile(student_id):
    """Update a profile by student_id"""
    try:
        data = request.get_json()

        # Get the profile from database
        from app.models.profile import Profile
        profile = get_db_session().query(Profile).filter(
            Profile.student_id == student_id,
            Profile.deleted_at.is_(None)
        ).first()

        if not profile:
            return jsonify({'error': 'Profile not found'}), 404
        
        # Update profile fields
        profile.candidate_name = data.get('candidate_name', profile.candidate_name)
        profile.email_id = data.get('email_id', profile.email_id)
        profile.contact_no = data.get('contact_no', profile.contact_no)
        profile.total_experience = data.get('total_experience', profile.total_experience)
        profile.relevant_experience = data.get('relevant_experience', profile.relevant_experience)
        profile.current_company = data.get('current_company', profile.current_company)
        profile.location = data.get('location', profile.location)
        profile.notice_period_days = data.get('notice_period_days', profile.notice_period_days)
        profile.ctc_current = data.get('ctc_current', profile.ctc_current)
        profile.ctc_expected = data.get('ctc_expected', profile.ctc_expected)
        profile.key_skills = data.get('key_skills', profile.key_skills)
        profile.education = data.get('education', profile.education)
        profile.source = data.get('source', profile.source)
        profile.resume_file_path = data.get('resume_file_path', profile.resume_file_path)
        profile.resume_file_name = data.get('resume_file_name', profile.resume_file_name)
        
        # Update timestamp
        profile.updated_at = datetime.utcnow()
        
        # Save to database
        get_db_session().commit()
        
        current_app.logger.info(f"Profile updated successfully: {student_id}")
        
        return jsonify({
            'success': True,
            'message': 'Profile updated successfully',
            'profile': profile.to_dict()
        })
        
    except Exception as e:
        current_app.logger.error(f"Error updating profile {student_id}: {str(e)}")
        get_db_session().rollback()
        return jsonify({'error': f'Failed to update profile: {str(e)}'}), 500

@api_bp.route('/profiles/<string:student_id>/resume', methods=['POST'])
def upload_resume(student_id):
    """Upload resume file for a profile"""
    try:
        current_app.logger.info(f"Resume upload request received for student_id: {student_id}")
        current_app.logger.info(f"Request files: {list(request.files.keys())}")
        current_app.logger.info(f"Request form data: {list(request.form.keys())}")

        # Check if file is present in request
        if 'resume' not in request.files:
            current_app.logger.error(f"No resume file in request. Available files: {list(request.files.keys())}")
            return jsonify({'error': 'No resume file provided'}), 400

        file = request.files['resume']

        # Check if file is selected
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400

        # Check file extension
        allowed_extensions = {'pdf', 'docx'}
        file_extension = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''

        if file_extension not in allowed_extensions:
            return jsonify({'error': 'Only PDF and DOCX files are allowed'}), 400

        # Get the profile from database
        from app.models.profile import Profile
        profile = get_db_session().query(Profile).filter(
            Profile.student_id == student_id,
            Profile.deleted_at.is_(None)
        ).first()

        if not profile:
            return jsonify({'error': 'Profile not found'}), 404
        
        # Generate unique filename
        import uuid
        import os
        from datetime import datetime
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        unique_id = str(uuid.uuid4())[:8]
        safe_filename = f"{student_id}_{timestamp}_{unique_id}.{file_extension}"
        
        # Create uploads directory if it doesn't exist
        upload_dir = os.path.join(current_app.root_path, 'uploads', 'resumes')
        os.makedirs(upload_dir, exist_ok=True)
        
        # Save file
        file_path = os.path.join(upload_dir, safe_filename)
        file.save(file_path)
        
        # Update profile with resume information
        profile.resume_file_path = safe_filename
        profile.resume_file_name = file.filename
        profile.updated_at = datetime.utcnow()
        
        # Save to database
        get_db_session().commit()
        
        current_app.logger.info(f"Resume uploaded successfully for profile: {student_id}")
        
        return jsonify({
            'success': True,
            'message': 'Resume uploaded successfully',
            'resume_file_name': file.filename,
            'resume_file_path': safe_filename
        })
        
    except Exception as e:
        current_app.logger.error(f"Error uploading resume for profile {student_id}: {str(e)}")
        get_db_session().rollback()
        return jsonify({'error': f'Failed to upload resume: {str(e)}'}), 500

@api_bp.route('/profiles/<string:student_id>/resume', methods=['GET'])
def download_resume(student_id):
    """Download resume file for a profile"""
    try:
        # Get the profile from database
        from app.models.profile import Profile
        profile = get_db_session().query(Profile).filter(
            Profile.student_id == student_id,
            Profile.deleted_at.is_(None)
        ).first()

        if not profile:
            return jsonify({'error': 'Profile not found'}), 404
        
        if not profile.resume_file_path:
            return jsonify({'error': 'No resume file found for this profile'}), 404
        
        # Construct file path
        import os
        file_path = os.path.join(current_app.root_path, 'uploads', 'resumes', profile.resume_file_path)
        
        if not os.path.exists(file_path):
            return jsonify({'error': 'Resume file not found on server'}), 404
        
        # Return file
        return send_from_directory(
            os.path.join(current_app.root_path, 'uploads', 'resumes'),
            profile.resume_file_path,
            as_attachment=True,
            download_name=profile.resume_file_name or f"resume_{student_id}.pdf"
        )
        
    except Exception as e:
        current_app.logger.error(f"Error downloading resume for profile {student_id}: {str(e)}")
        return jsonify({'error': f'Failed to download resume: {str(e)}'}), 500

@api_bp.route('/convert-docx-to-html', methods=['GET'])
def convert_docx_to_html():
    """Convert DOCX file to HTML for viewing"""
    try:
        file_path = request.args.get('path')
        if not file_path:
            return jsonify({'success': False, 'error': 'File path is required'}), 400

        # Security: Validate file path is within uploads directory
        import os
        from werkzeug.utils import secure_filename
        
        # Normalize the path to prevent directory traversal
        normalized_path = os.path.normpath(file_path)
        uploads_dir = os.path.join(current_app.root_path, 'uploads')
        
        # Handle both absolute and relative paths
        if os.path.isabs(normalized_path):
            # It's an absolute path, check if it's within uploads directory
            if not normalized_path.startswith(uploads_dir):
                return jsonify({'success': False, 'error': 'Access denied'}), 403
        else:
            # It's a relative path, construct the full path
            normalized_path = os.path.join(uploads_dir, normalized_path)
        
        # Check if file exists
        if not os.path.exists(normalized_path):
            return jsonify({'success': False, 'error': 'File not found'}), 404
        
        # Check file extension
        file_extension = os.path.splitext(normalized_path)[1].lower()
        if file_extension != '.docx':
            return jsonify({'success': False, 'error': 'Only DOCX files are supported'}), 400
        
        # Convert DOCX to HTML
        from docx import Document
        import re
        
        try:
            doc = Document(normalized_path)
            html_content = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Convert paragraph to HTML
                    text = paragraph.text.strip()
                    # Escape HTML characters
                    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    
                    # Check paragraph style for headings and lists
                    style_name = paragraph.style.name if paragraph.style else ""
                    is_heading = style_name.startswith('Heading') or style_name.startswith('Title')
                    is_list_item = style_name.startswith('List') or style_name.startswith('Bullet')
                    is_numbered_list = style_name.startswith('Numbered') or 'Number' in style_name
                    
                    # Check if paragraph has any runs with formatting
                    has_formatting = False
                    for run in paragraph.runs:
                        if run.bold or run.italic or run.underline:
                            has_formatting = True
                            break
                    
                    if is_heading:
                        # Determine heading level
                        heading_level = 1
                        if 'Heading 1' in style_name:
                            heading_level = 1
                        elif 'Heading 2' in style_name:
                            heading_level = 2
                        elif 'Heading 3' in style_name:
                            heading_level = 3
                        elif 'Heading 4' in style_name:
                            heading_level = 4
                        elif 'Heading 5' in style_name:
                            heading_level = 5
                        elif 'Heading 6' in style_name:
                            heading_level = 6
                        else:
                            heading_level = 2  # Default to h2 for other headings
                        
                        html_content += f"<h{heading_level} class='font-semibold text-gray-900 mt-6 mb-3'>{text}</h{heading_level}>\n"
                    elif is_list_item:
                        # Simple bullet list item
                        html_content += f"<li class='mb-2 ml-4'>{text}</li>\n"
                    elif is_numbered_list:
                        # Simple numbered list item
                        html_content += f"<li class='mb-2 ml-4'>{text}</li>\n"
                    elif has_formatting:
                        # Process runs with formatting
                        formatted_text = ""
                        for run in paragraph.runs:
                            run_text = run.text
                            if run_text:
                                # Escape HTML characters
                                run_text = run_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                
                                # Apply formatting
                                if run.bold:
                                    run_text = f"<strong>{run_text}</strong>"
                                if run.italic:
                                    run_text = f"<em>{run_text}</em>"
                                if run.underline:
                                    run_text = f"<u>{run_text}</u>"
                                
                                formatted_text += run_text
                        
                        html_content += f"<p class='mb-3 leading-relaxed'>{formatted_text}</p>\n"
                    else:
                        # Simple paragraph
                        html_content += f"<p class='mb-3 leading-relaxed'>{text}</p>\n"
            
            # Process tables if any
            for table in doc.tables:
                html_content += "<table class='border-collapse border border-gray-300 w-full my-4'>\n"
                for i, row in enumerate(table.rows):
                    html_content += "<tr>\n"
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        cell_text = cell_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        tag = "th" if i == 0 else "td"
                        html_content += f"<{tag} class='border border-gray-300 px-3 py-2'>{cell_text}</{tag}>\n"
                    html_content += "</tr>\n"
                html_content += "</table>\n"
            
            # Process any additional content in the document
            # Check for headers and footers
            for section in doc.sections:
                # Process header content
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text = paragraph.text.strip()
                            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            html_content += f"<div class='header-content text-xs text-gray-500 border-b pb-2 mb-4'>{text}</div>\n"
                
                # Process footer content
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            text = paragraph.text.strip()
                            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            html_content += f"<div class='footer-content text-xs text-gray-500 border-t pt-2 mt-4'>{text}</div>\n"
            
            # If no content was extracted, return a message
            if not html_content.strip():
                html_content = "<p>No readable content found in the document.</p>"
            
            return jsonify({
                'success': True,
                'html_content': html_content
            })
            
        except Exception as e:
            current_app.logger.error(f"Error converting DOCX to HTML: {str(e)}")
            return jsonify({'success': False, 'error': 'Failed to convert document'}), 500
            
    except Exception as e:
        current_app.logger.error(f"Error in convert_docx_to_html: {str(e)}")
        return jsonify({'success': False, 'error': 'Internal server error'}), 500

@api_bp.route('/download-file', methods=['GET'])
def download_file():
    """Download file with security validation"""
    try:
        file_path = request.args.get('path')
        if not file_path:
            return jsonify({'error': 'File path is required'}), 400

        # Security: Validate file path is within uploads directory
        import os
        from werkzeug.utils import secure_filename
        
        # Normalize the path to prevent directory traversal
        normalized_path = os.path.normpath(file_path)
        uploads_dir = os.path.join(current_app.root_path, 'uploads')
        
        # Handle both absolute and relative paths
        if os.path.isabs(normalized_path):
            # It's an absolute path, check if it's within uploads directory
            if not normalized_path.startswith(uploads_dir):
                return jsonify({'error': 'Access denied'}), 403
        else:
            # It's a relative path, construct the full path
            normalized_path = os.path.join(uploads_dir, normalized_path)
        
        # Check if file exists
        if not os.path.exists(normalized_path):
            return jsonify({'error': 'File not found'}), 404
        
        # Get the directory and filename
        file_dir = os.path.dirname(normalized_path)
        filename = os.path.basename(normalized_path)
        
        # Determine MIME type based on file extension
        file_extension = os.path.splitext(filename)[1].lower()
        mime_types = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.txt': 'text/plain',
            '.html': 'text/html',
            '.htm': 'text/html'
        }
        
        mimetype = mime_types.get(file_extension, 'application/octet-stream')
        
        return send_from_directory(
            file_dir,
            filename,
            as_attachment=True,
            mimetype=mimetype
        )
        
    except Exception as e:
        current_app.logger.error(f"Error in download_file: {str(e)}")
        return jsonify({'error': 'Internal server error'}), 500

@api_bp.route('/view-file', methods=['GET'])
def view_file():
    """View file in browser (for PDFs)"""
    try:
        file_path = request.args.get('path')
        if not file_path:
            return jsonify({'error': 'File path is required'}), 400

        # Security: Validate file path is within uploads directory
        import os
        
        # Normalize the path to prevent directory traversal
        normalized_path = os.path.normpath(file_path)
        uploads_dir = os.path.join(current_app.root_path, 'uploads')
        
        # Handle both absolute and relative paths
        if os.path.isabs(normalized_path):
            # It's an absolute path, check if it's within uploads directory
            if not normalized_path.startswith(uploads_dir):
                return jsonify({'error': 'Access denied'}), 403
        else:
            # It's a relative path, construct the full path
            normalized_path = os.path.join(uploads_dir, normalized_path)
        
        # Check if file exists
        if not os.path.exists(normalized_path):
            return jsonify({'error': 'File not found'}), 404
        
        # Get the directory and filename
        file_dir = os.path.dirname(normalized_path)
        filename = os.path.basename(normalized_path)
        
        # Only allow PDF files for viewing
        file_extension = os.path.splitext(filename)[1].lower()
        if file_extension != '.pdf':
            return jsonify({'error': 'Only PDF files can be viewed'}), 400
        
        return send_from_directory(
            file_dir,
            filename,
            mimetype='application/pdf'
        )
        
    except Exception as e:
        current_app.logger.error(f"Error in view_file: {str(e)}")
        return jsonify({'error': 'Internal server error'}), 500

@api_bp.route('/recruiter/login', methods=['POST'])
def recruiter_login():
    """Login for existing recruiter with domain isolation"""
    try:
        data = request.get_json()
        username = data.get('username')
        password = data.get('password')
        
        if not username or not password:
            return jsonify({
                'status': 'error',
                'message': 'Username and password are required'
            }), 400
        
        # Get domain from custom header or fallback to detection
        domain = request.headers.get('X-Original-Domain')
        if not domain:
            domain = request.headers.get('X-Domain')
        
        # Ensure domain database isolation before authentication
        if not ensure_domain_isolation():
            current_app.logger.error(f"Failed to ensure domain database isolation for domain: {domain}")
            return jsonify({
                'status': 'error',
                'message': 'Database not available for this domain'
            }), 503
        
        # Get domain-specific database session
        from flask import g
        if not hasattr(g, 'db_session') or g.db_session is None:
            current_app.logger.error("No database session available for domain")
            return jsonify({
                'status': 'error',
                'message': 'Database session not available for this domain'
            }), 503
        
        # Find user by username in domain-specific database
        user = g.db_session.query(User).filter_by(username=username).first()
        
        if not user or not user.check_password(password):
            return jsonify({
                'status': 'error',
                'message': 'Invalid username or password'
            }), 401
        
        # Create JWT token
        access_token = create_access_token(identity=username)
        
        return jsonify({
            'status': 'success',
            'message': 'Login successful',
            'access_token': access_token,
            'user': user.to_dict()
        })
        
    except Exception as e:
        current_app.logger.error(f"Error in recruiter login: {str(e)}")
        return jsonify({
            'status': 'error',
            'message': 'Failed to login'
        }), 500

@api_bp.route('/send-email', methods=['POST'])
def send_email():
    """Send email from recruiter workflow"""
    try:
        data = request.get_json()
        recipient = data.get('to')
        subject = data.get('subject')
        body = data.get('body')
        request_id = data.get('request_id')
        
        if not all([recipient, subject, body]):
            return jsonify({'error': 'Missing required fields: to, subject, body'}), 400
        
        # Here you would integrate with your email service
        # For now, we'll simulate sending and log the email
        current_app.logger.info(f"Email sent to {recipient} for request {request_id}")
        current_app.logger.info(f"Subject: {subject}")
        current_app.logger.info(f"Body: {body}")
        
        # Use Microsoft Graph API to send email
        email_processor = EmailProcessor()
        result = email_processor.send_email(recipient, subject, body, request_id)
        
        if result['success']:
            return jsonify(result)
        else:
            return jsonify({'error': result['error']}), 500
        
    except Exception as e:
        current_app.logger.error(f"Error sending email: {str(e)}")
        return jsonify({'error': f'Failed to send email: {str(e)}'}), 500

@api_bp.route('/users/recruiters', methods=['GET'])
def get_recruiters():
    """Get all recruiter users for admin assignment"""
    try:
        current_app.logger.info("Fetching all recruiters")
        recruiters = get_db_session().query(User).filter_by(role='recruiter').order_by(User.username).all()
        current_app.logger.info(f"Found {len(recruiters)} recruiters")
        
        result = [recruiter.to_dict() for recruiter in recruiters]
        
        current_app.logger.info("Successfully converted recruiters to JSON")
        return jsonify(result)
    except Exception as e:
        current_app.logger.error(f"Error fetching recruiters data: {str(e)}")
        return jsonify({'error': 'Failed to fetch recruiters data'}), 500

@api_bp.route('/teams-meeting', methods=['POST'])
def create_teams_meeting():
    """Create a Microsoft Teams meeting and return the meeting link"""
    try:
        data = request.get_json()
        
        # Required fields
        subject = data.get('subject')
        start_time = data.get('start_time')  # ISO format: "2024-01-15T10:00:00Z"
        end_time = data.get('end_time')      # ISO format: "2024-01-15T11:00:00Z"
        attendees = data.get('attendees', [])  # List of email addresses
        request_id = data.get('request_id')
        meeting_type = data.get('meeting_type', 'interview')  # interview, round1, round2
        
        if not all([subject, start_time, end_time]):
            return jsonify({
                'success': False,
                'error': 'Missing required fields: subject, start_time, end_time'
            }), 400
        
        # Use existing EmailProcessor to create meeting
        email_processor = EmailProcessor()
        
        # Get candidate_id from request data or attendees
        candidate_id = data.get('candidate_id')
        if not candidate_id and attendees and len(attendees) > 0:
            # Try to find the candidate profile to get student_id
            from app.models.profile import Profile
            candidate_profile = get_db_session().query(Profile).filter_by(email_id=attendees[0]).first()
            if candidate_profile:
                candidate_id = candidate_profile.student_id
        
        result = email_processor.create_teams_meeting(
            subject=subject,
            start_time=start_time,
            end_time=end_time,
            attendees=attendees,
            request_id=request_id,
            meeting_type=meeting_type,
            candidate_id=candidate_id
        )
        
        if result['success']:
            # Persist meeting to DB (minimal replacement for in-memory/local storage)
            try:
                if candidate_id:
                    # Parse ISO times
                    from datetime import datetime
                    parsed_start = None
                    parsed_end = None
                    try:
                        parsed_start = datetime.fromisoformat(start_time.replace('Z', '+00:00')) if start_time else None
                        parsed_end = datetime.fromisoformat(end_time.replace('Z', '+00:00')) if end_time else None
                    except Exception as parse_err:
                        current_app.logger.warning(f"Failed to parse meeting times: {parse_err}")

                    Meeting.upsert(
                        request_id=request_id or '',
                        candidate_id=candidate_id,
                        round_type=meeting_type or 'interview_scheduled',
                        meet_link=result.get('teams_meeting_link', ''),
                        start_time=parsed_start,
                        end_time=parsed_end,
                        timezone='UTC',
                        subject=result.get('subject', subject),
                    )
                else:
                    current_app.logger.warning("No candidate_id provided; skipping meeting DB save")
            except Exception as e:
                current_app.logger.error(f"Error saving meeting to DB: {str(e)}")

            # Keep existing in-memory store as secondary fallback (optional)
            if candidate_id:
                try:
                    from app.services.calendar_service import CalendarService
                    calendar_service = CalendarService()
                    calendar_service.store_meeting_info(
                        request_id=request_id,
                        candidate_id=candidate_id,
                        round_type=meeting_type,
                        meet_link=result.get('teams_meeting_link', ''),
                        start_time=start_time,
                        end_time=end_time,
                        subject=result.get('subject', subject)
                    )
                except Exception as e:
                    current_app.logger.error(f"Error storing meeting info in memory: {str(e)}")

            return jsonify(result)
        else:
            return jsonify({'error': result['error']}), 500
            
    except Exception as e:
        current_app.logger.error(f"Error creating Teams meeting: {str(e)}")
        return jsonify({'error': f'Failed to create Teams meeting: {str(e)}'}), 500

@api_bp.route('/send-interview-email/<string:request_id>', methods=['POST'])
def send_interview_email(request_id):
    """Send interview email with Teams meeting link"""
    try:
        # Handle both JSON and multipart form data
        if request.content_type and 'multipart/form-data' in request.content_type:
            # Handle multipart form data (with file attachments)
            recipient_email = request.form.get('recipient_email')
            recipient_name = request.form.get('recipient_name', 'Candidate')
            cc_email = request.form.get('cc_email', '')
            subject = request.form.get('subject')
            body = request.form.get('body')
            teams_meeting_link = request.form.get('teams_meeting_link')
            meeting_details = request.form.get('meeting_details', '{}')
            selected_profiles = request.form.get('selected_profiles', '[]')
            interview_step = request.form.get('interview_step', 'interview_scheduled')
            
            # Parse JSON strings
            import json
            try:
                meeting_details = json.loads(meeting_details) if meeting_details else {}
                selected_profiles = json.loads(selected_profiles) if selected_profiles else []
            except json.JSONDecodeError:
                return jsonify({'error': 'Invalid JSON format in form data'}), 400
            
            # Get file attachments
            attachments = request.files.getlist('attachments')
        else:
            # Handle JSON data (backward compatibility)
            data = request.get_json()
            recipient_email = data.get('recipient_email')
            recipient_name = data.get('recipient_name', 'Candidate')
            cc_email = data.get('cc_email', '')
            subject = data.get('subject')
            body = data.get('body')
            teams_meeting_link = data.get('teams_meeting_link')
            meeting_details = data.get('meeting_details', {})
            selected_profiles = data.get('selected_profiles', [])
            interview_step = data.get('interview_step', 'interview_scheduled')  # interview_scheduled, round1, round2
            attachments = []
        
        if not all([recipient_email, subject, body]):
            return jsonify({'error': 'Missing required fields: recipient_email, subject, body'}), 400
        
        # Get requirement details
        requirement = get_db_session().query(Requirement).filter_by(request_id=request_id).first()
        if not requirement:
            return jsonify({'error': 'Requirement not found'}), 404
        
        # Send email using email service
        email_processor = EmailProcessor()
        result = email_processor.send_interview_email(
            to_email=recipient_email,
            subject=subject,
            body=body,
            cc_email=cc_email,
            teams_meeting_link=teams_meeting_link,
            meeting_details=meeting_details,
            request_id=request_id,
            requirement=requirement,
            interview_step=interview_step,
            attachments=attachments
        )
        
        if result['success']:
            return jsonify(result)
        else:
            return jsonify({'error': result['error']}), 500
            
    except Exception as e:
        current_app.logger.error(f"Error sending interview email: {str(e)}")
        return jsonify({'error': f'Failed to send interview email: {str(e)}'}), 500


@api_bp.route('/meet-links/<string:request_id>', methods=['GET'])
def get_meet_links(request_id):
    """Get Teams meeting links, preferring DB over calendar/memory"""
    try:
        round_type = request.args.get('round_type')
        candidate_id = request.args.get('candidate_id')

        # 1) Try database first (minimal new behavior)
        if candidate_id:
            meeting = Meeting.get_for_candidate(
                request_id=request_id,
                candidate_id=candidate_id,
                round_type=round_type or 'interview_scheduled',
            )
            if meeting:
                return jsonify({'success': True, 'meet_info': meeting})
        else:
            meet_links = Meeting.get_for_request(
                request_id=request_id, round_type=round_type
            )
            if meet_links:
                return jsonify({'success': True, 'meet_links': meet_links})

        # 2) Fallback to calendar/memory if DB is empty
        from app.services.calendar_service import CalendarService, _global_meeting_storage
        calendar_service = CalendarService()

        if candidate_id:
            # Get meet link for specific candidate
            meet_info = calendar_service.get_meet_link_for_candidate(
                request_id=request_id,
                candidate_id=candidate_id,
                round_type=round_type or 'interview_scheduled'
            )
            
            if meet_info:
                # Save fallback-found meeting to DB for persistence
                try:
                    from datetime import datetime
                    Meeting.upsert(
                        request_id=request_id,
                        candidate_id=candidate_id,
                        round_type=round_type or 'interview_scheduled',
                        meet_link=meet_info.get('meet_link', ''),
                        start_time=(
                            datetime.fromisoformat(meet_info['start_time'].replace('Z', '+00:00'))
                            if meet_info.get('start_time')
                            else None
                        ),
                        end_time=(
                            datetime.fromisoformat(meet_info['end_time'].replace('Z', '+00:00'))
                            if meet_info.get('end_time')
                            else None
                        ),
                        timezone=meet_info.get('timezone', 'UTC'),
                        subject=meet_info.get('subject'),
                    )
                except Exception as e:
                    current_app.logger.error(f"Failed to persist fallback meeting: {e}")
                return jsonify({
                    'success': True,
                    'meet_info': meet_info
                })
            else:
                return jsonify({
                    'success': False,
                    'message': 'No meeting found for this candidate'
                })
        else:
            # Get all meet links for the request
            meet_links = calendar_service.get_meet_links_for_request(
                request_id=request_id,
                round_type=round_type
            )
            
            # Persist any found meetings
            try:
                from datetime import datetime
                for cand_id, info in meet_links.items():
                    Meeting.upsert(
                        request_id=request_id,
                        candidate_id=cand_id,
                        round_type=info.get('round_type') or (round_type or 'interview_scheduled'),
                        meet_link=info.get('meet_link', ''),
                        start_time=(
                            datetime.fromisoformat(info['start_time'].replace('Z', '+00:00'))
                            if info.get('start_time')
                            else None
                        ),
                        end_time=(
                            datetime.fromisoformat(info['end_time'].replace('Z', '+00:00'))
                            if info.get('end_time')
                            else None
                        ),
                        timezone=info.get('timezone', 'UTC'),
                        subject=info.get('subject'),
                    )
            except Exception as e:
                current_app.logger.error(f"Failed to persist fallback meetings: {e}")

            return jsonify({'success': True, 'meet_links': meet_links})
            
    except Exception as e:
        current_app.logger.error(f"Error getting meet links: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Failed to get meet links: {str(e)}'
        }), 500

@api_bp.route('/get-enum-values', methods=['GET'])
def get_enum_values():
    """Get all enum values for a specific enum type"""
    try:
        from flask import request
        from app.database import db
        from sqlalchemy import text
        
        enum_type = request.args.get('enum_type')
        
        if not enum_type:
            return jsonify({'error': 'enum_type parameter is required'}), 400
        
        # Validate enum type
        valid_enum_types = {
            'company': 'companyenum',
            'department': 'departmentenum', 
            'shift': 'shiftenum',
            'job_type': 'jobtypeenum',
            'priority': 'priorityenum'
        }
        
        if enum_type not in valid_enum_types:
            return jsonify({'error': 'Invalid enum type'}), 400
        
        db_enum_type = valid_enum_types[enum_type]
        
        # Get all enum values
        query = text("""
            SELECT enumlabel 
            FROM pg_enum 
            WHERE enumtypid = (
                SELECT oid FROM pg_type WHERE typname = :enum_type
            ) 
            ORDER BY enumlabel
        """)
        
        result = get_db_session().execute(query, {'enum_type': db_enum_type}).fetchall()
        
        enum_values = [row[0] for row in result]
        
        return jsonify({
            'success': True,
            'enum_type': enum_type,
            'values': enum_values
        })
        
    except Exception as e:
        current_app.logger.error(f"Error getting enum values: {str(e)}")
        return jsonify({'error': f'Failed to get enum values: {str(e)}'}), 500

@api_bp.route('/add-enum-value', methods=['POST'])
def add_enum_value():
    """Add a new value to an existing enum type in the database"""
    try:
        from flask import request
        from app.database import db
        from sqlalchemy import text
        
        data = request.get_json()
        current_app.logger.info(f"Received data: {data}")
        
        if not data:
            return jsonify({'error': 'No JSON data provided'}), 400
            
        enum_type = data.get('enum_type')
        new_value = data.get('new_value')
        
        current_app.logger.info(f"enum_type: {enum_type}, new_value: {new_value}")
        
        if not enum_type or not new_value:
            return jsonify({'error': f'enum_type and new_value are required. Got: enum_type={enum_type}, new_value={new_value}'}), 400
        
        # Validate enum type
        valid_enum_types = {
            'company': 'companyenum',
            'department': 'departmentenum', 
            'shift': 'shiftenum',
            'job_type': 'jobtypeenum',
            'priority': 'priorityenum'
        }
        
        if enum_type not in valid_enum_types:
            return jsonify({'error': 'Invalid enum type'}), 400
        
        db_enum_type = valid_enum_types[enum_type]
        
        # Sanitize the new value (replace spaces with underscores and make lowercase for consistency)
        sanitized_value = new_value.replace(' ', '_').lower()
        
        # Check if the value already exists
        check_query = text("""
            SELECT enumlabel 
            FROM pg_enum 
            WHERE enumtypid = (
                SELECT oid FROM pg_type WHERE typname = :enum_type
            ) AND enumlabel = :enum_value
        """)
        
        result = get_db_session().execute(check_query, {
            'enum_type': db_enum_type,
            'enum_value': sanitized_value
        }).fetchone()
        
        if result:
            return jsonify({'error': f'Value "{new_value}" already exists in {enum_type} enum'}), 400
        
        # Add the new enum value
        alter_query = text(f"ALTER TYPE {db_enum_type} ADD VALUE :enum_value")
        
        try:
            get_db_session().execute(alter_query, {'enum_value': sanitized_value})
            get_db_session().commit()
            
            return jsonify({
                'success': True,
                'message': f'Successfully added "{new_value}" to {enum_type} enum',
                'added_value': sanitized_value,
                'display_value': new_value  # Return the original value for display
            })
            
        except Exception as db_error:
            get_db_session().rollback()
            current_app.logger.error(f"Database error adding enum value: {str(db_error)}")
            return jsonify({
                'error': f'Database error: {str(db_error)}'
            }), 500
            
    except Exception as e:
        current_app.logger.error(f"Error adding enum value: {str(e)}")
        return jsonify({'error': f'Failed to add enum value: {str(e)}'}), 500

@api_bp.route('/recruiter-activity', methods=['GET'])
def get_recruiter_activity():
    """Get daily recruiter activity data"""
    try:
        # Get query parameters
        days = request.args.get('days', default=7, type=int)  # Default to last 7 days
        date_from = request.args.get('date_from', type=str)
        date_to = request.args.get('date_to', type=str)
        
        # Calculate date range
        if date_from and date_to:
            start_date = datetime.strptime(date_from, '%Y-%m-%d').date()
            end_date = datetime.strptime(date_to, '%Y-%m-%d').date()
        else:
            end_date = datetime.now().date()
            start_date = end_date - timedelta(days=days-1)
        
        # Get all recruiters
        recruiters = get_db_session().query(User).filter_by(role='recruiter').all()
        recruiter_usernames = [r.username for r in recruiters]
        
        # Get daily activity data
        daily_activity = []
        current_date = start_date
        
        while current_date <= end_date:
            date_str = current_date.strftime('%Y-%m-%d')
            
            # Get profiles submitted on this date, include requirement_id for attribution
            profiles_submitted = get_db_session().query(
                Profile.profile_id,
                Profile.requirement_id,
                Profile.created_at
            ).filter(
                func.date(Profile.created_at) == current_date
            ).all()
            
            # Get all requirements and their assigned recruiters
            all_requirements = get_db_session().query(
                Requirement.requirement_id,
                Requirement.user_id
            ).all()
            
            # Create a mapping of requirement_id to assigned recruiters
            requirement_recruiters = {}
            for req in all_requirements:
                if req.user_id:
                    # Get the user (recruiter) assigned to this requirement
                    user = get_db_session().query(User).filter_by(user_id=req.user_id).first()
                    if user:
                        requirement_recruiters[str(req.requirement_id)] = [user.username]
            
            # Create a mapping of profile_id to requirement_id using the new schema
            profile_to_requirement = {}
            for profile in profiles_submitted:
                # profile tuple includes requirement_id already from the query
                if getattr(profile, 'requirement_id', None):
                    profile_to_requirement[str(profile.profile_id)] = str(profile.requirement_id)
            
            # Count profiles per recruiter for this date
            recruiter_counts = {recruiter: 0 for recruiter in recruiter_usernames}
            for profile in profiles_submitted:
                req_id = profile_to_requirement.get(str(profile.profile_id))
                if req_id and req_id in requirement_recruiters:
                    assigned_recruiters = requirement_recruiters[req_id]
                    for recruiter in assigned_recruiters:
                        if recruiter in recruiter_usernames:
                            recruiter_counts[recruiter] += 1
                            break  # Count only once per profile
            
            # Get active recruiters (consider active if submitted at least 1 profile)
            active_recruiters = [
                {
                    'username': recruiter,
                    'profiles_submitted': count,
                    'is_active': count > 0
                }
                for recruiter, count in recruiter_counts.items()
            ]
            
            # Sort by profiles submitted (descending)
            active_recruiters.sort(key=lambda x: x['profiles_submitted'], reverse=True)
            
            # Calculate totals
            total_profiles = sum(count for count in recruiter_counts.values())
            active_recruiter_count = sum(1 for count in recruiter_counts.values() if count > 0)
            
            daily_activity.append({
                'date': date_str,
                'total_profiles_submitted': total_profiles,
                'active_recruiters_count': active_recruiter_count,
                'recruiters': active_recruiters
            })
            
            current_date += timedelta(days=1)
        
        # Get overall statistics
        total_recruiters = len(recruiter_usernames)
        
        # Get today's activity
        today = datetime.now().date()
        today_activity = next((day for day in daily_activity if day['date'] == today.strftime('%Y-%m-%d')), None)
        
        # Get this week's activity
        week_start = today - timedelta(days=today.weekday())
        week_activity = [
            day for day in daily_activity 
            if datetime.strptime(day['date'], '%Y-%m-%d').date() >= week_start
        ]
        
        # Calculate weekly totals
        weekly_total_profiles = sum(day['total_profiles_submitted'] for day in week_activity)
        weekly_active_recruiters = set(
            recruiter['username']
            for day in week_activity
            for recruiter in day['recruiters']
            if recruiter['is_active']
        )
        
        # Company-wise recruiter performance using new schema (no legacy tracker table)
        company_performance_query = """
        SELECT 
            r.company_name,
            u.username as recruiter_name,
            COUNT(*) as total_profiles_submitted,
            COUNT(CASE WHEN p.status = 'onboarded' THEN 1 END) as onboarded_profiles,
            CASE 
                WHEN COUNT(*) > 0 THEN 
                    ROUND((COUNT(CASE WHEN p.status = 'onboarded' THEN 1 END)::numeric / COUNT(*)) * 100, 2)
                ELSE 0.0 
            END as success_rate
        FROM profiles p
        JOIN requirements r ON p.requirement_id = r.requirement_id
        JOIN users u ON p.created_by_recruiter = u.user_id
        WHERE p.created_by_recruiter IS NOT NULL 
        AND u.role = 'recruiter'
        AND LOWER(u.username) != 'admin'
        GROUP BY r.company_name, u.username
        ORDER BY r.company_name, onboarded_profiles DESC
        """
        
        company_performance_result = get_db_session().execute(company_performance_query)
        
        # Group by company
        company_data = {}
        for row in company_performance_result:
            company_name = row.company_name or 'Unknown Company'
            if company_name not in company_data:
                company_data[company_name] = []
            
            company_data[company_name].append({
                'recruiter_name': row.recruiter_name,
                'total_profiles_submitted': row.total_profiles_submitted,
                'onboarded_profiles': row.onboarded_profiles,
                'success_rate': row.success_rate
            })
        
        return jsonify({
            'success': True,
            'data': {
                'date_range': {
                    'start_date': start_date.strftime('%Y-%m-%d'),
                    'end_date': end_date.strftime('%Y-%m-%d')
                },
                'overall_stats': {
                    'total_recruiters': total_recruiters,
                    'today_active_recruiters': today_activity['active_recruiters_count'] if today_activity else 0,
                    'today_profiles_submitted': today_activity['total_profiles_submitted'] if today_activity else 0,
                    'weekly_active_recruiters': len(weekly_active_recruiters),
                    'weekly_profiles_submitted': weekly_total_profiles
                },
                'daily_activity': daily_activity,
                'recruiter_list': recruiter_usernames,
                'company_performance': company_data
            }
        })
        
    except Exception as e:
        current_app.logger.error(f"Error fetching recruiter activity: {str(e)}")
        return jsonify({
            'success': False,
            'error': 'Failed to fetch recruiter activity data'
        }), 500


@api_bp.route('/requirements-activity', methods=['GET'])
@require_domain_auth
def get_requirements_activity():
    """Get requirements activity data"""
    try:
        # Get today's date and week start
        today = datetime.now().date()
        week_start = today - timedelta(days=today.weekday())
        
        # Get total RFH count (all requirements with requirement_id)
        total_rfh = get_db_session().query(Requirement).filter(
            Requirement.requirement_id.isnot(None)
        ).count()
        
        # Get today's active requirements (requirements that recruiters worked on today)
        # This includes requirements where profiles were submitted today
        today_active_requirements = get_db_session().query(
            Requirement.requirement_id
        ).join(
            Profile, Requirement.requirement_id == Profile.requirement_id
        ).filter(
            func.date(Profile.created_at) == today
        ).distinct().count()
        
        # Get this week's active requirements (requirements that recruiters worked on this week)
        weekly_active_requirements = get_db_session().query(
            Requirement.requirement_id
        ).join(
            Profile, Requirement.requirement_id == Profile.requirement_id
        ).filter(
            func.date(Profile.created_at) >= week_start
        ).distinct().count()
        
        return jsonify({
            'success': True,
            'data': {
                'total_rfh': total_rfh,
                'today_active_requirements': today_active_requirements,
                'weekly_active_requirements': weekly_active_requirements
            }
        })
        
    except Exception as e:
        current_app.logger.error(f"Error fetching requirements activity: {str(e)}")
        return jsonify({
            'success': False,
            'error': 'Failed to fetch requirements activity data'
        }), 500


@api_bp.route('/send-inactive-recruiter-notifications', methods=['POST'])
def send_inactive_recruiter_notifications():
    """Send email notifications to inactive recruiters (admin only)"""
    try:
        # Check if user is admin
        from flask import session
        user_id = session.get('user_id')
        if not user_id:
            return jsonify({
                'success': False,
                'error': 'Authentication required'
            }), 401
        
        user = get_db_session().query(User).filter_by(id=str(user_id)).first()
        if not user or user.role != 'admin':
            return jsonify({
                'success': False,
                'error': 'Admin access required'
            }), 403
        
        # Initialize notification service
        notification_service = RecruiterNotificationService()
        
        # Send manual notifications (bypasses time/day restrictions)
        result = notification_service.send_manual_notifications()
        
        return jsonify(result)
        
    except Exception as e:
        current_app.logger.error(f"Error sending inactive recruiter notifications: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Failed to send notifications: {str(e)}'
        }), 500


@api_bp.route('/get-inactive-recruiters', methods=['GET'])
def get_inactive_recruiters():
    """Get list of inactive recruiters for today (admin only)"""
    try:
        # Check if user is admin
        from flask import session
        user_id = session.get('user_id')
        if not user_id:
            return jsonify({
                'success': False,
                'error': 'Authentication required'
            }), 401
        
        user = get_db_session().query(User).filter_by(id=str(user_id)).first()
        if not user or user.role != 'admin':
            return jsonify({
                'success': False,
                'error': 'Admin access required'
            }), 403
        
        # Initialize notification service
        notification_service = RecruiterNotificationService()
        
        # Get inactive recruiters
        inactive_recruiters = notification_service.get_inactive_recruiters_for_today()
        
        return jsonify({
            'success': True,
            'data': {
                'inactive_recruiters': inactive_recruiters,
                'count': len(inactive_recruiters),
                'timestamp': datetime.now().isoformat()
            }
        })
        
    except Exception as e:
        current_app.logger.error(f"Error getting inactive recruiters: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Failed to get inactive recruiters: {str(e)}'
        }), 500


@api_bp.route('/scheduler/status', methods=['GET'])
def get_scheduler_status_endpoint():
    """Get scheduler status and job information (admin only)"""
    try:
        # Check if user is admin
        from flask import session
        user_id = session.get('user_id')
        if not user_id:
            return jsonify({
                'success': False,
                'error': 'Authentication required'
            }), 401
        
        user = get_db_session().query(User).filter_by(id=str(user_id)).first()
        if not user or user.role != 'admin':
            return jsonify({
                'success': False,
                'error': 'Admin access required'
            }), 403
        
        # Get scheduler status
        status = get_scheduler_status()
        
        return jsonify({
            'success': True,
            'data': status
        })
        
    except Exception as e:
        current_app.logger.error(f"Error getting scheduler status: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Failed to get scheduler status: {str(e)}'
        }), 500


@api_bp.route('/scheduler/pause', methods=['POST'])
def pause_scheduler_endpoint():
    """Pause the scheduler (admin only)"""
    try:
        # Check if user is admin
        from flask import session
        user_id = session.get('user_id')
        if not user_id:
            return jsonify({
                'success': False,
                'error': 'Authentication required'
            }), 401
        
        user = get_db_session().query(User).filter_by(id=str(user_id)).first()
        if not user or user.role != 'admin':
            return jsonify({
                'success': False,
                'error': 'Admin access required'
            }), 403
        
        # Pause scheduler
        result = pause_scheduler()
        
        return jsonify(result)
        
    except Exception as e:
        current_app.logger.error(f"Error pausing scheduler: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Failed to pause scheduler: {str(e)}'
        }), 500


@api_bp.route('/scheduler/resume', methods=['POST'])
def resume_scheduler_endpoint():
    """Resume the scheduler (admin only)"""
    try:
        # Check if user is admin
        from flask import session
        user_id = session.get('user_id')
        if not user_id:
            return jsonify({
                'success': False,
                'error': 'Authentication required'
            }), 401
        
        user = get_db_session().query(User).filter_by(id=str(user_id)).first()
        if not user or user.role != 'admin':
            return jsonify({
                'success': False,
                'error': 'Admin access required'
            }), 403
        
        # Resume scheduler
        result = resume_scheduler()
        
        return jsonify(result)
        
    except Exception as e:
        current_app.logger.error(f"Error resuming scheduler: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Failed to resume scheduler: {str(e)}'
        }), 500


@api_bp.route('/scheduler/run-job/<string:job_id>', methods=['POST'])
def run_job_manually_endpoint(job_id):
    """Run a specific job manually (admin only)"""
    try:
        # Check if user is admin
        from flask import session
        user_id = session.get('user_id')
        if not user_id:
            return jsonify({
                'success': False,
                'error': 'Authentication required'
            }), 401
        
        user = get_db_session().query(User).filter_by(id=str(user_id)).first()
        if not user or user.role != 'admin':
            return jsonify({
                'success': False,
                'error': 'Admin access required'
            }), 403
        
        # Run job manually
        result = run_job_manually(job_id)
        
        return jsonify(result)
        
    except Exception as e:
        current_app.logger.error(f"Error running job manually: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Failed to run job manually: {str(e)}'
        }), 500

@api_bp.route('/test-domain-isolation', methods=['GET'])
def test_domain_isolation():
    """Test endpoint to verify domain isolation is working"""
    try:
        # Get domain from headers
        domain = request.headers.get('X-Original-Domain') or request.headers.get('X-Domain')
        
        # Get database session info
        session = get_db_session()
        session_id = id(session)
        
        # Get some basic counts from the database
        user_count = session.query(User).count()
        profile_count = session.query(Profile).count()
        requirement_count = session.query(Requirement).count()
        
        # Get database connection info if available
        db_info = {}
        try:
            if hasattr(session, 'bind') and hasattr(session.bind, 'url'):
                db_info['database_url'] = str(session.bind.url)
        except:
            pass
        
        return jsonify({
            'success': True,
            'domain': domain,
            'session_id': session_id,
            'database_counts': {
                'users': user_count,
                'profiles': profile_count,
                'requirements': requirement_count
            },
            'database_info': db_info,
            'is_domain_session': hasattr(g, 'db_session') and g.db_session is not None,
            'message': 'Domain isolation test completed'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e),
            'message': 'Domain isolation test failed'
        }), 500


 