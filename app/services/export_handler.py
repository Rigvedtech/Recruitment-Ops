import os
import pandas as pd
from datetime import datetime
from typing import List, Dict, Any

class ExportHandler:
    def __init__(self):
        # Use absolute path to ensure correct directory
        import os
        current_dir = os.path.dirname(os.path.abspath(__file__))
        self.export_dir = os.path.join(current_dir, '..', 'uploads', 'exports')
        os.makedirs(self.export_dir, exist_ok=True)

    def export_data(self, emails: List[Dict[str, Any]]) -> str:
        """Export email data to Excel"""
        try:
            # Process emails for export
            processed_data = []
            for email in emails:
                processed_email = {
                    'Subject': email.get('subject', ''),
                    'Sender': email.get('sender', {}).get('emailAddress', {}).get('address', '') if isinstance(email.get('sender'), dict) else email.get('sender', ''),
                    'Received Date': email.get('receivedDateTime', ''),
                    'Attachments': len(email.get('attachments', [])),
                    'Student Name': email.get('student_info', {}).get('name', ''),
                    'Experience': email.get('student_info', {}).get('total_experience', ''),
                    'Current Company': email.get('student_info', {}).get('current_company', ''),
                    'Current CTC': email.get('student_info', {}).get('current_ctc', ''),
                    'Expected CTC': email.get('student_info', {}).get('expected_ctc', ''),
                    'Notice Period': email.get('student_info', {}).get('notice_period', ''),
                    'Location': email.get('student_info', {}).get('location', ''),
                    'Education': email.get('student_info', {}).get('education', ''),
                    'Skills': email.get('student_info', {}).get('skills', ''),
                }
                processed_data.append(processed_email)
            
            # Create DataFrame
            df = pd.DataFrame(processed_data)
            
            # Generate timestamp for filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'email_data_{timestamp}.xlsx'
            file_path = os.path.join(self.export_dir, filename)
            
            # Export to Excel with formatting
            with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Email Data')
                worksheet = writer.sheets['Email Data']
                
                # Auto-adjust column widths
                for column in worksheet.columns:
                    max_length = 0
                    column = [cell for cell in column]
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = (max_length + 2)
                    worksheet.column_dimensions[column[0].column_letter].width = adjusted_width
            
            return file_path
        except Exception as e:
            print(f"Error exporting data: {str(e)}")
            raise e

    def export_profiles(self, profiles: List[Dict[str, Any]], request_id: str, hiring_manager_name: str = "Hiring Manager") -> str:
        """Export profile data as Word document"""
        try:
            # Import required libraries for Word document creation
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.shared import OxmlElement, qn
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml
            
            # Generate timestamp for filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'profiles_{request_id}_{timestamp}.docx'
            file_path = os.path.join(self.export_dir, filename)
            
            # Debug logging
            print(f"Export directory: {self.export_dir}")
            print(f"Export file path: {file_path}")
            print(f"Directory exists: {os.path.exists(self.export_dir)}")
            
            # Create Word document
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Add greeting
            greeting_para = doc.add_paragraph()
            greeting_run = greeting_para.add_run(f"Dear {hiring_manager_name},")
            greeting_run.font.size = Pt(12)
            
            # Add spacing
            doc.add_paragraph()
            
            # Add introduction
            intro_para = doc.add_paragraph()
            intro_run = intro_para.add_run("Please find below the candidate profiles for your review:")
            intro_run.font.size = Pt(12)
            
            # Add spacing
            doc.add_paragraph()
            
            # Create table
            table = doc.add_table(rows=1, cols=10)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'Table Grid'
            
            # Set table headers
            headers = [
                'Candidate Name', 'Contact', 'Total Experience (Years)', 
                'Relevant Experience (Years)', 'Location', 'CTC Current', 
                'CTC Expected', 'Key Skills', 'Source', 'Education'
            ]
            
            header_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].font.bold = True
                header_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
            
            # Add profile rows
            for profile in profiles:
                row_cells = table.add_row().cells
                
                candidate_name = profile.get('candidate_name', '')
                contact = profile.get('contact_no', '') or profile.get('email_id', profile.get('candidate_email', ''))
                total_exp = profile.get('total_experience', 0)
                relevant_exp = profile.get('relevant_experience', 0)
                location = profile.get('location', '')
                ctc_current = profile.get('ctc_current', 0)
                ctc_expected = profile.get('ctc_expected', 0)
                key_skills = profile.get('key_skills', '')
                source = profile.get('source', '')
                education = profile.get('education', '')
                
                # Fill row data
                row_cells[0].text = str(candidate_name)
                row_cells[1].text = str(contact)
                row_cells[2].text = str(total_exp)
                row_cells[3].text = str(relevant_exp)
                row_cells[4].text = str(location)
                row_cells[5].text = str(ctc_current)
                row_cells[6].text = str(ctc_expected)
                row_cells[7].text = str(key_skills)
                row_cells[8].text = str(source)
                row_cells[9].text = str(education)
                
                # Set font size for all cells in the row
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
            
            # Add spacing
            doc.add_paragraph()
            
            # Add signature
            signature_para = doc.add_paragraph()
            signature_run = signature_para.add_run("Regards,\nRecruitment Ops")
            signature_run.font.size = Pt(12)
            
            # Save the document
            doc.save(file_path)
            
            # Debug logging after file creation
            print(f"File created successfully: {file_path}")
            print(f"File exists after creation: {os.path.exists(file_path)}")
            print(f"File size: {os.path.getsize(file_path) if os.path.exists(file_path) else 'N/A'}")
            
            return file_path
        except Exception as e:
            print(f"Error exporting profiles: {str(e)}")
            raise e

    def generate_email_content(self, profiles: List[Dict[str, Any]], request_id: str, requirement: Any, recipient_name: str = "Hiring Manager", selected_columns: List[str] = None) -> str:
        """Generate HTML email content with simplified template"""
        try:
            # Define column mappings
            column_mappings = {
                'candidate_name': 'Candidate Name',
                'contact_no': 'Contact Number',
                'email_id': 'Email Address',
                'total_experience': 'Total Exp (Yrs)',
                'relevant_experience': 'Relevant Exp (Yrs)',
                'current_company': 'Current Company',
                'location': 'Location',
                'ctc_current': 'Current CTC (LPA)',
                'ctc_expected': 'Expected CTC (LPA)',
                'notice_period_days': 'Notice Period (Days)',
                'key_skills': 'Key Skills',
                'education': 'Education',
                'source': 'Source'
            }
            
            # Use selected columns or default to all columns
            if selected_columns is None:
                selected_columns = ['candidate_name', 'contact_no', 'total_experience', 'relevant_experience', 'current_company', 'location', 'ctc_current', 'ctc_expected', 'key_skills']
            
            # Ensure candidate_name is always included
            if 'candidate_name' not in selected_columns:
                selected_columns.insert(0, 'candidate_name')
            
            # Create HTML content with minimal template
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 20px; }}
                    table {{ width: 100%; border-collapse: collapse; margin-bottom: 20px; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; font-size: 12px; }}
                    th {{ background-color: #f8f9fa; font-weight: bold; }}
                </style>
            </head>
            <body>
                <p>Dear {recipient_name},</p>
                
                <p>Please find below the candidate profiles for your review:</p>
                
                <table>
                    <thead>
                        <tr>
                            {''.join([f'<th>{column_mappings.get(col, col)}</th>' for col in selected_columns])}
                        </tr>
                    </thead>
                    <tbody>
            """
            
            # Add profile rows
            for profile in profiles:
                html_content += "<tr>"
                for col in selected_columns:
                    if col == 'contact_no':
                        # Special handling for contact - combine contact_no and email
                        contact = profile.get('contact_no', '') or profile.get('email_id', profile.get('candidate_email', ''))
                        html_content += f"<td>{contact}</td>"
                    else:
                        value = profile.get(col, '')
                        # Format numeric values
                        if col in ['total_experience', 'relevant_experience', 'ctc_current', 'ctc_expected']:
                            value = str(value) if value else '0'
                        elif col == 'notice_period_days':
                            value = f"{value} days" if value else 'N/A'
                        else:
                            value = str(value) if value else 'N/A'
                        html_content += f"<td>{value}</td>"
                html_content += "</tr>"
            
            html_content += f"""
                    </tbody>
                </table>
                
                <p>Regards,<br>
                Recruitment Ops</p>
            </body>
            </html>
            """
            
            return html_content
            
        except Exception as e:
            print(f"Error generating email content: {str(e)}")
            raise 