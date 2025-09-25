import os
from datetime import datetime, timedelta
import msal
import requests
import json
from typing import Dict, List, Optional, Any, Union
import re
from bs4 import BeautifulSoup
import html2text
import pandas as pd
from io import StringIO
import chardet
import html
from email import policy
from email.parser import Parser
import unicodedata
from tabulate import tabulate
from pandas.core.frame import DataFrame
from html2markdown import convert as convert_to_markdown
from sqlalchemy.orm import Session, scoped_session
from app.database import db
from config import Config
from app.models.profile import Profile

class EmailService:
    def __init__(self):
        self.config = Config
        self.client_id = Config.MS_CLIENT_ID
        self.client_secret = Config.MS_CLIENT_SECRET
        self.authority = Config.MS_AUTHORITY
        self.scope = Config.MS_SCOPE
        self.user_email = Config.MS_USER_EMAIL
        self.html2text = html2text.HTML2Text()
        self.html2text.ignore_links = True
        self.html2text.body_width = 0
        self.html2text.unicode_snob = True
        self.html2text.ul_item_mark = '-'
        self.html2text.emphasis_mark = '*'
        self.html2text.strong_mark = '**'
        
        # Key-value pair patterns with improved regex and validation
        self.field_patterns = {
            'name': r'(?:Name|Candidate Name|Full Name|Candidate|Resource Name)\s*[:-]\s*([A-Za-z\s\.]{2,50})',
            'total_experience': r'(?:Total Experience|Total Exp|Experience|Total Exp|Relevant Exp|Years of Experience|Exp\(Yrs\)|Years|Yrs)\s*[:-]\s*(\d+(?:\.\d+)?)\s*(?:years|yrs|Yrs|Years)?',
            'current_company': r'(?:Current Company|Company|Organization|Current Organization|Present Company|Current Employer|Current Organization)\s*[:-]\s*([A-Za-z0-9\s\.,&-]{2,50})',
            'current_ctc': r'(?:Current CTC|CTC \(Current\)|Current Package|Current Salary|CTC|Current Annual CTC)\s*[:-]\s*(?:₹|Rs\.?|INR)?\s*(\d+(?:\.\d+)?)\s*(?:LPA|Lakhs?|L|Lacs|lac|p\.a\.|per annum)',
            'expected_ctc': r'(?:Expected CTC|CTC \(Expected\)|Expected Package|Expected Salary|ECTC|Expected Annual CTC)\s*[:-]\s*(?:₹|Rs\.?|INR)?\s*(\d+(?:\.\d+)?)\s*(?:LPA|Lakhs?|L|Lacs|lac|p\.a\.|per annum)',
            'notice_period': r'(?:Notice Period|NP|Notice|Current Notice Period|Notice Period \(Days\))\s*[:-]\s*(\d+)\s*(?:days?|Days?|months?|Months?|week|weeks)',
            'location': r'(?:Location|City|Current Location|Present Location|Base Location|Work Location)\s*[:-]\s*([A-Za-z\s\.,\-]{2,50}?)(?:\s*<|\s*$)',
            'education': r'(?:Education|Qualification|Highest Qualification|Educational Background|Educational Qualification)\s*[:-]\s*([A-Za-z0-9\s\.,\(\)-]{2,100})',
            'skills': r'(?:Skills|Key Skills|Technical Skills|Primary Skills|Core Skills|Technology|Technologies|Tech Stack)\s*[:-]\s*([A-Za-z0-9\s\.,+#-]{2,200})',
            'source': r'(?:Source|Reference|Referred By|Referral|Reference Source)\s*[:-]\s*([A-Za-z0-9\s\.,@-]{2,50})'
        }
        
        # Column headers mapping for tabular data
        self.table_headers = {
            'Name': 'name',
            'Candidate Name': 'name',
            'Full Name': 'name',
            'Total Experience': 'total_experience',
            'Total Exp': 'total_experience',
            'Experience': 'total_experience',
            'Current Company': 'current_company',
            'Company': 'current_company',
            'Organization': 'current_company',
            'Current CTC': 'current_ctc',
            'CTC (Current)': 'current_ctc',
            'Current Package': 'current_ctc',
            'Expected CTC': 'expected_ctc',
            'CTC (Expected)': 'expected_ctc',
            'Expected Package': 'expected_ctc',
            'Notice Period': 'notice_period',
            'NP': 'notice_period',
            'Notice': 'notice_period',
            'Location': 'location',
            'City': 'location',
            'Current Location': 'location',
            'Education': 'education',
            'Qualification': 'education',
            'Highest Qualification': 'education',
            'Skills': 'skills',
            'Key Skills': 'skills',
            'Technical Skills': 'skills',
            'Source': 'source',
            'Reference': 'source',
            'Referred By': 'source'
        }

    def _clean_text(self, text: str) -> str:
        """Clean and normalize HTML text while preserving structure"""
        try:
            # First try to detect the encoding
            encoding_result = chardet.detect(text.encode())
            if encoding_result['encoding']:
                text = text.encode().decode(encoding_result['encoding'])
            
            # Initialize html2text
            h = html2text.HTML2Text()
            h.ignore_links = False
            h.body_width = 0  # Don't wrap text
            h.unicode_snob = True  # Use Unicode characters
            h.ul_item_mark = '-'  # Use - for unordered lists
            h.emphasis_mark = '*'  # Use * for emphasis
            h.strong_mark = '**'  # Use ** for strong
            
            # Convert HTML to markdown
            markdown_text = h.handle(text)
            
            # Clean up the markdown text
            lines = markdown_text.split('\n')
            cleaned_lines = []
            for line in lines:
                line = line.strip()
                if line:  # Only keep non-empty lines
                    cleaned_lines.append(line)
            
            # Join lines with proper spacing
            cleaned_text = '\n\n'.join(cleaned_lines)
            
            return cleaned_text
        except Exception as e:
            print(f"Error cleaning text: {str(e)}")
            return text  # Return original text if cleaning fails

    def _validate_extracted_value(self, field: str, value: str) -> str:
        """Validate and clean extracted values"""
        if not value or not isinstance(value, str):
            return ''
        
        value = value.strip()
        if not value:
            return ''
        
        try:
            if field == 'name':
                # Name should be 2-50 chars, only letters, spaces, dots, and hyphens
                if not re.match(r'^[A-Za-z\s.-]{2,50}$', value):
                    return ''
                # Capitalize each word
                return ' '.join(word.capitalize() for word in value.split())
                
            elif field == 'total_experience':
                # Convert various experience formats to years
                value = value.lower()
                # Remove any text after numbers (e.g., "5+ years" -> "5")
                value = re.sub(r'(\d+(?:\.\d+)?)\s*(?:years?|yrs?|\+|\s*$).*', r'\1', value)
                try:
                    exp = float(value)
                    if 0 <= exp <= 50:  # Reasonable experience range
                        return f"{exp:.1f}"
                except ValueError:
                    pass
                return ''
                
            elif field == 'current_company':
                # Company name should be 2-100 chars, letters, numbers, spaces, and basic punctuation
                if not re.match(r'^[A-Za-z0-9\s.&@,-]{2,100}$', value):
                    return ''
                return value.strip()
                
            elif field in ['current_ctc', 'expected_ctc']:
                # Convert various salary formats to a standard format
                value = value.lower()
                # Remove currency symbols and 'lpa'/'per annum'
                value = re.sub(r'[₹$€£]|\s*lpa|\s*per\s*annum', '', value)
                # Extract the first number (handle ranges like "10-12")
                match = re.search(r'(\d+(?:\.\d+)?)', value)
                if match:
                    try:
                        ctc = float(match.group(1))
                        if 0 < ctc <= 100:  # Reasonable CTC range in lakhs
                            return f"{ctc:.2f}"
                    except ValueError:
                        pass
                return ''
                
            elif field == 'notice_period':
                # Convert notice period to days
                value = value.lower()
                if 'immediate' in value or 'asap' in value:
                    return '0'
                
                # Extract number and unit
                match = re.search(r'(\d+)\s*(day|week|month)', value)
                if match:
                    num = int(match.group(1))
                    unit = match.group(2)
                    if unit.startswith('week'):
                        num *= 7
                    elif unit.startswith('month'):
                        num *= 30
                    if 0 <= num <= 180:  # Reasonable notice period range
                        return str(num)
                return ''
                
            elif field == 'location':
                # Location should be 2-50 chars, letters, spaces, and basic punctuation
                if not re.match(r'^[A-Za-z\s,.-]{2,50}$', value):
                    return ''
                # Capitalize each word
                return ' '.join(word.capitalize() for word in value.split())
                
            elif field == 'education':
                # Education should contain common degree abbreviations or full names
                value = value.upper()
                degrees = ['BE', 'BTECH', 'ME', 'MTECH', 'BCA', 'MCA', 'BSC', 'MSC', 'MBA', 'PHD']
                if not any(degree in value.replace('.', '') for degree in degrees):
                    return ''
                return value
                
            elif field == 'skills':
                # Skills should be comma-separated technical terms
                skills = [s.strip() for s in value.split(',')]
                valid_skills = []
                for skill in skills:
                    # Basic validation for technical skills
                    if re.match(r'^[A-Za-z0-9#.+\-]{2,20}$', skill):
                        valid_skills.append(skill)
                return ', '.join(valid_skills) if valid_skills else ''
                
            elif field == 'source':
                # Source should not look like a phone number or email
                if re.search(r'\d{10}|@', value):  # Skip if it looks like phone/email
                    return ''
                if not re.match(r'^[A-Za-z\s,.-]{2,50}$', value):
                    return ''
                return value.strip()
                
        except Exception as e:
            print(f"Error validating {field}: {str(e)}", flush=True)
            return ''
        
        return value.strip()

    def _extract_table_data(self, text: str) -> dict[str, str]:
        """Extract data from tabular format"""
        print("\n=== Attempting to extract table data ===", flush=True)
        
        data = {
            'name': '',
            'total_experience': '',
            'current_company': '',
            'current_ctc': '',
            'expected_ctc': '',
            'notice_period': '',
            'location': '',
            'education': '',
            'skills': '',
            'source': ''
        }
        
        try:
            # Try to find tables in the text
            tables: List[DataFrame] = []
            
            # Method 1: Parse HTML tables with multiple parsers
            try:
                # Try different HTML parsers in order of reliability
                parsers = ['html5lib', 'lxml', 'html.parser']
                for parser in parsers:
                    try:
                        soup = BeautifulSoup(text, parser)
                        html_tables = soup.find_all('table')
                        print(f"Found {len(html_tables)} HTML tables using {parser}", flush=True)
                        
                        if html_tables:
                            for table in html_tables:
                                try:
                                    # Convert HTML table to DataFrame
                                    rows = []
                                    headers: List[str] = []
                                    
                                    # Get headers from th or first tr
                                    header_row = table.find('tr')
                                    if header_row:
                                        headers = [th.get_text(strip=True) for th in header_row.find_all(['th', 'td'])]
                                    
                                    # Get data rows
                                    data_rows = table.find_all('tr')[1:] if headers else table.find_all('tr')
                                    for row in data_rows:
                                        cols = [td.get_text(strip=True) for td in row.find_all('td')]
                                        if cols and (len(cols) == len(headers) if headers else True):
                                            rows.append(cols)
                                    
                                    if rows:
                                        if not headers and rows:
                                            # If no headers found, try to detect them from first row
                                            potential_headers = rows[0]
                                            if any(header.lower() in self.table_headers for header in potential_headers):
                                                headers = potential_headers
                                                rows = rows[1:]
                                        
                                        # Create DataFrame with proper typing
                                        if headers:
                                            df = pd.DataFrame(data=rows)
                                            df.columns = pd.Index(headers)
                                        else:
                                            df = pd.DataFrame(data=rows)
                                        
                                        tables.append(df)
                                        print(f"Found table using {parser} parser", flush=True)
                                except Exception as e:
                                    print(f"Error parsing HTML table with {parser}: {str(e)}", flush=True)
                                    continue
                        
                            if tables:  # If we found tables, break the parser loop
                                break
                    except Exception as e:
                        print(f"Error with {parser} parser: {str(e)}", flush=True)
                    continue
            except Exception as e:
                print(f"HTML table parsing failed: {str(e)}", flush=True)
            
            # Method 2: Convert markdown tables to HTML and parse
            if not tables:
                try:
                    # Look for markdown table patterns
                    markdown_patterns = [
                        # Standard markdown tables
                        r'(\|[^\n]+\|\n\|[-:\|\s]+\|\n(?:\|[^\n]+\|\n)+)',
                        # Simple pipe tables
                        r'(\|[^\n]+\|\n(?:\|[^\n]+\|\n)+)',
                        # Space-aligned tables
                        r'([^\n]+\|[^\n]+\n[-:\|\s]+\n(?:[^\n]+\|[^\n]+\n)+)'
                    ]
                    
                    for pattern in markdown_patterns:
                        markdown_tables = re.finditer(pattern, text, re.MULTILINE)
                        for match in markdown_tables:
                            markdown_table = match.group(1)
                            
                            # Clean up the table
                            lines = [line.strip() for line in markdown_table.split('\n') if line.strip()]
                            if len(lines) >= 2:  # Need at least header and one data row
                                try:
                                    # Convert to DataFrame
                                    df = pd.read_csv(StringIO('\n'.join(lines)), sep='|', skipinitialspace=True, engine='python')
                                    df = df.dropna(axis=1, how='all')  # Remove empty columns
                                    if not df.empty:
                                        tables.append(df)
                                        print("Found table using markdown pattern", flush=True)
                                except Exception as e:
                                    print(f"Error parsing markdown table: {str(e)}", flush=True)
                                    continue
                except Exception as e:
                    print(f"Markdown conversion failed: {str(e)}", flush=True)
            
            # Method 3: Look for structured text patterns
            if not tables:
                try:
                    # Split text into lines and look for potential table structures
                    lines = text.split('\n')
                    potential_table_start = -1
                    current_table_lines = []
                    
                    for i, line in enumerate(lines):
                        # Look for lines that might be table headers
                        if any(header.lower() in line.lower() for header in self.table_headers.keys()):
                            # Check if next line could be data
                            if i + 1 < len(lines) and lines[i + 1].strip():
                                potential_table_start = i
                                current_table_lines = [line]
                        elif potential_table_start >= 0:
                            # Continue collecting lines that look like data
                            if line.strip() and len(line.split()) >= 2:
                                current_table_lines.append(line)
                            else:
                                # End of potential table
                                if len(current_table_lines) >= 2:
                                    try:
                                        # Try to parse as space-separated table
                                        df = pd.read_csv(StringIO('\n'.join(current_table_lines)), 
                                                       delim_whitespace=True,
                                                       engine='python')
                                        if not df.empty:
                                            tables.append(df)
                                            print("Found table using structured text pattern", flush=True)
                                    except Exception:
                                        pass  # If parsing fails, just continue
                                potential_table_start = -1
                                current_table_lines = []
                except Exception as e:
                    print(f"Structured text parsing failed: {str(e)}", flush=True)
            
            # Process all found tables
            for df in tables:
                print("\nAnalyzing table:", flush=True)
                print(tabulate(df.head(), headers='keys', tablefmt='pipe'), flush=True)
                
                # Clean column names
                df.columns = df.columns.str.strip()
                
                # Look for known headers in the table
                for header in self.table_headers.keys():
                    matching_cols = [col for col in df.columns if header.lower() in col.lower()]
                    for col in matching_cols:
                        try:
                            # Get first non-null value
                            values = df[col].dropna()
                            if not values.empty:
                                value = str(values.iloc[0]).strip()
                                field_name = self.table_headers[header]
                                if not data[field_name]:  # Only update if not already set
                                    validated_value = self._validate_extracted_value(field_name, value)
                                    if validated_value:
                                        data[field_name] = validated_value
                                        print(f"Found {field_name}: {validated_value}", flush=True)
                        except Exception as e:
                            print(f"Error processing column {col}: {str(e)}", flush=True)
                
                # If we found any data, try to find more in the same table
                if any(data.values()):
                    continue
            
        except Exception as e:
            print(f"Error in table extraction: {str(e)}", flush=True)
        
        print(f"\nExtracted table data: {data}", flush=True)
        return data

    def _extract_key_value_data(self, text: str) -> dict[str, str]:
        """Extract data using regex patterns for key-value format"""
        print("\n=== Attempting to extract key-value data ===", flush=True)
        
        data = {
            'name': '',
            'total_experience': '',
            'current_company': '',
            'current_ctc': '',
            'expected_ctc': '',
            'notice_period': '',
            'location': '',
            'education': '',
            'skills': '',
            'source': ''
        }
        
        # Split text into lines for line-by-line processing
        lines = text.split('\n')
        
        for line in lines:
            # Skip empty lines
            if not line.strip():
                continue
            
            # Try each pattern on the current line
            for field, pattern in self.field_patterns.items():
                if data[field]:  # Skip if we already have this field
                    continue
                    
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    value = match.group(1).strip()
                    validated_value = self._validate_extracted_value(field, value)
                    if validated_value:
                        data[field] = validated_value
                        print(f"Found {field}: {validated_value}", flush=True)
                    break
        
        print(f"\nExtracted key-value data: {data}", flush=True)
        return data

    def _extract_student_data(self, text: str) -> dict[str, str]:
        """Extract student data using multiple methods"""
        # First clean the text
        cleaned_text = self._clean_text(text)
        
        # Try table format first
        data = self._extract_table_data(cleaned_text)
        
        # If no data found, try key-value format
        if not any(data.values()):
            data = self._extract_key_value_data(cleaned_text)
        
        return data

    def _extract_student_info(self, text: str) -> dict:
        """Extract student information from email text"""
        try:
            # Convert HTML to text first
            clean_text = self._clean_text(text)
            
            # Initialize student info dictionary
            student_info = {
                'name': '',
                'total_exp': '',
                'relevant_exp': '',
                'current_company': '',
                'current_ctc': '',
                'expected_ctc': '',
                'notice_period': '',
                'location': '',
                'education': '',
                'key_skills': '',
                'source': '',
                'resume_attached': False
            }
            
            # Split text into lines for processing
            lines = clean_text.split('\n')
            
            # Process each line
            for line in lines:
                line = line.strip()
                
                # Skip empty lines
                if not line:
                    continue
                
                # Look for table-like data (pipe separated)
                if '|' in line:
                    cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                    if len(cells) >= 2:
                        # Try to identify the type of data based on the content
                        for cell in cells:
                            if 'yrs' in cell.lower() or 'years' in cell.lower():
                                if not student_info['total_exp']:
                                    student_info['total_exp'] = cell
                            elif 'lpa' in cell.lower() or '₹' in cell:
                                if not student_info['current_ctc']:
                                    student_info['current_ctc'] = cell
                                elif not student_info['expected_ctc']:
                                    student_info['expected_ctc'] = cell
                            elif 'days' in cell.lower() or 'months' in cell.lower():
                                student_info['notice_period'] = cell
                            elif any(edu in cell.lower() for edu in ['b.tech', 'b.e', 'bachelor', 'm.tech', 'mca']):
                                student_info['education'] = cell
                            elif any(tech in cell.lower() for tech in ['java', 'python', 'react', 'node', 'angular']):
                                student_info['key_skills'] = cell
                            elif any(source in cell.lower() for source in ['naukri', 'linkedin', 'indeed']):
                                student_info['source'] = cell
                            elif any(loc in cell.lower() for loc in ['bangalore', 'pune', 'hyderabad', 'mumbai', 'delhi']):
                                student_info['location'] = cell
                            elif any(company in cell.lower() for company in ['tcs', 'infosys', 'wipro', 'ibm', 'accenture']):
                                student_info['current_company'] = cell
                
                # Look for key-value pairs
                if ':' in line:
                    key, value = line.split(':', 1)
                    key = key.strip().lower()
                    value = value.strip()
                    
                    if 'name' in key and not student_info['name']:
                        student_info['name'] = value
                    elif 'experience' in key and not student_info['total_exp']:
                        student_info['total_exp'] = value
                    elif 'company' in key and not student_info['current_company']:
                        student_info['current_company'] = value
                    elif 'ctc' in key and not student_info['current_ctc']:
                        student_info['current_ctc'] = value
                    elif 'expected' in key and not student_info['expected_ctc']:
                        student_info['expected_ctc'] = value
                    elif 'notice' in key and not student_info['notice_period']:
                        student_info['notice_period'] = value
                    elif 'location' in key and not student_info['location']:
                        student_info['location'] = value
                    elif 'education' in key and not student_info['education']:
                        student_info['education'] = value
                    elif 'skills' in key and not student_info['key_skills']:
                        student_info['key_skills'] = value
                    elif 'source' in key and not student_info['source']:
                        student_info['source'] = value
            
            # Check if resume is attached
            student_info['resume_attached'] = 'resume' in clean_text.lower() or 'attached' in clean_text.lower()
            
            return student_info
        except Exception as e:
            print(f"Error extracting student info: {str(e)}")
            return {}

    def _process_email(self, email_data: dict) -> dict:
        """Process a single email and extract relevant information"""
        try:
            # Extract basic email information
            email_id = email_data.get('id', '')
            subject = email_data.get('subject', '')
            sender = email_data.get('sender', {}).get('emailAddress', {}).get('address', '')
            received_date = email_data.get('receivedDateTime', '')
            body = email_data.get('body', {}).get('content', '')
            
            # Clean and extract information from the email body
            clean_body = self._clean_text(body)
            student_info = self._extract_student_info(body)
            
            # Save to database if we have student info
            if student_info:
                profile_data = {
                    'email_id': email_id,
                    'name': student_info.get('name', ''),
                    'total_exp': student_info.get('total_exp', ''),
                    'relevant_exp': student_info.get('relevant_exp', ''),
                    'current_company': student_info.get('current_company', ''),
                    'current_ctc': student_info.get('current_ctc', ''),
                    'expected_ctc': student_info.get('expected_ctc', ''),
                    'notice_period': student_info.get('notice_period', ''),
                    'location': student_info.get('location', ''),
                    'education': student_info.get('education', ''),
                    'key_skills': student_info.get('key_skills', ''),
                    'source': student_info.get('source', ''),
                    'resume_attached': student_info.get('resume_attached', False),
                    'email_subject': subject,
                    'email_sender': sender,
                    'received_date': received_date
                }
                
                # Create a new Profile instance
                profile = Profile(**profile_data)
                db.add(profile)
                db.commit()
            
            return {
                'id': email_id,
                'subject': subject,
                'sender': sender,
                'receivedDateTime': received_date,
                'body_preview': clean_body,
                'student_info': student_info
            }
        except Exception as e:
            print(f"Error processing email: {str(e)}")
            if 'db' in locals():
                db.rollback()
            return {
                'id': email_id if 'email_id' in locals() else '',
                'subject': subject if 'subject' in locals() else '',
                'sender': sender if 'sender' in locals() else '',
                'receivedDateTime': received_date if 'received_date' in locals() else '',
                'body_preview': 'Error processing email content',
                'error': str(e)
            }

    def send_email(self, to_email: str, subject: str, html_content: str, profiles=None, request_id=None, requirement=None, recipient_name="Hiring Manager", cc_email="", attachments=None) -> bool:
        """Send email using Microsoft Graph API with fallback to console output"""
        try:
            # First try Microsoft Graph API
            access_token = self._get_access_token()
            if access_token:
                # Prepare email message
                message = {
                    "message": {
                        "subject": subject,
                        "body": {
                            "contentType": "HTML",
                            "content": html_content
                        },
                        "toRecipients": [
                            {
                                "emailAddress": {
                                    "address": to_email
                                }
                            }
                        ]
                    }
                }
                
                # Add Cc recipients if provided
                if cc_email and cc_email.strip():
                    # Handle multiple Cc emails (comma-separated)
                    cc_emails = [email.strip() for email in cc_email.split(',') if email.strip()]
                    if cc_emails:
                        message["message"]["ccRecipients"] = [
                            {
                                "emailAddress": {
                                    "address": email
                                }
                            }
                            for email in cc_emails
                        ]
                
                # Add attachments if provided
                if attachments:
                    import base64
                    message["message"]["attachments"] = []
                    for attachment in attachments:
                        try:
                            # Read file content
                            attachment.seek(0)  # Ensure we're at the beginning of the file
                            file_content = attachment.read()
                            encoded_content = base64.b64encode(file_content).decode('utf-8')
                            
                            # Determine content type
                            content_type = "application/octet-stream"  # Default
                            if attachment.filename.lower().endswith('.pdf'):
                                content_type = "application/pdf"
                            elif attachment.filename.lower().endswith('.docx'):
                                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            
                            message["message"]["attachments"].append({
                                "@odata.type": "#microsoft.graph.fileAttachment",
                                "name": attachment.filename,
                                "contentType": content_type,
                                "contentBytes": encoded_content
                            })
                        except Exception as e:
                            print(f"Error processing attachment {attachment.filename}: {str(e)}")
                            continue
                
                # Send email using Microsoft Graph API
                url = f"https://graph.microsoft.com/v1.0/users/{self.user_email}/sendMail"
                headers = {
                    'Authorization': f'Bearer {access_token}',
                    'Content-Type': 'application/json'
                }
                
                response = requests.post(url, headers=headers, json=message)
                
                if response.status_code == 202:  # 202 Accepted is the expected response
                    print(f"Email sent successfully to {to_email}")
                    return True
                else:
                    print(f"Failed to send email via Graph API. Status: {response.status_code}, Response: {response.text}")
                    # Fall through to fallback method
            else:
                print("Failed to get access token for Microsoft Graph API, using fallback method")
                
        except Exception as e:
            print(f"Error with Microsoft Graph API: {str(e)}, using fallback method")
        
        # Fallback: Save email as Word document and show in console
        try:
            import os
            from datetime import datetime
            
            # Create emails directory if it doesn't exist
            emails_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'uploads', 'emails')
            os.makedirs(emails_dir, exist_ok=True)
            
            # Generate filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"email_{timestamp}_{to_email.replace('@', '_at_').replace('.', '_')}.docx"
            filepath = os.path.join(emails_dir, filename)
            
            # Create Word document with profiles data
            try:
                if profiles and request_id and requirement:
                    # Use ExportHandler to create proper Word document
                    from app.services.export_handler import ExportHandler
                    export_handler = ExportHandler()
                    
                    # Create Word document with the same template as export
                    word_file_path = export_handler.export_profiles(profiles, request_id, recipient_name)
                    
                    # Copy the generated file to our emails directory
                    import shutil
                    shutil.copy2(word_file_path, filepath)
                    
                else:
                    # Fallback: Create simple Word document
                    from docx import Document
                    from docx.shared import Inches, Pt
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    
                    doc = Document()
                    
                    # Add header
                    header_para = doc.add_paragraph()
                    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    header_run = header_para.add_run("Candidate Profiles Report")
                    header_run.font.size = Pt(16)
                    header_run.font.bold = True
                    
                    # Add email details
                    doc.add_paragraph()
                    info_para = doc.add_paragraph()
                    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    info_run = info_para.add_run(f"To: {to_email}")
                    info_run.font.size = Pt(10)
                    info_run.font.italic = True
                    
                    subject_para = doc.add_paragraph()
                    subject_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    subject_run = subject_para.add_run(f"Subject: {subject}")
                    subject_run.font.size = Pt(10)
                    subject_run.font.italic = True
                    
                    # Add spacing
                    doc.add_paragraph()
                    
                    # Add email content
                    content_para = doc.add_paragraph()
                    content_run = content_para.add_run("Email content has been generated and saved as a Word document.")
                    content_run.font.size = Pt(12)
                    
                    # Save the document
                    doc.save(filepath)
                    
            except ImportError:
                # If python-docx is not available, save as text
                with open(filepath.replace('.docx', '.txt'), 'w', encoding='utf-8') as f:
                    f.write(f"Email saved at: {datetime.now().isoformat()}\n")
                    f.write(f"To: {to_email}\n")
                    f.write(f"Subject: {subject}\n")
                    f.write("Email content has been generated.\n")
            
            print(f"\n{'='*80}")
            print(f"EMAIL SENT (Fallback Mode)")
            print(f"{'='*80}")
            print(f"To: {to_email}")
            print(f"Subject: {subject}")
            print(f"Saved to: {filepath}")
            print(f"{'='*80}")
            print("Email content preview:")
            print(f"{'='*80}")
            
            # Show a preview of the email content
            lines = html_content.split('\n')
            for i, line in enumerate(lines[:20]):  # Show first 20 lines
                print(line)
            if len(lines) > 20:
                print("... (truncated)")
            
            print(f"{'='*80}")
            print("To configure Microsoft Graph API for real email sending:")
            print("1. Set up Azure App Registration")
            print("2. Configure environment variables in .env file")
            print("3. Grant appropriate permissions to the app")
            print(f"{'='*80}")
            
            return True
            
        except Exception as e:
            print(f"Error in fallback email method: {str(e)}")
            return False

    def _get_access_token(self) -> str:
        """Get Microsoft Graph API access token"""
        try:
            print(f"Attempting to get access token with:")
            print(f"Client ID: {self.client_id}")
            print(f"Authority: {self.authority}")
            print(f"Scope: {self.scope}")
            print(f"User Email: {self.user_email}")
            
            # Create MSAL app
            app = msal.ConfidentialClientApplication(
                self.client_id,
                authority=self.authority,
                client_credential=self.client_secret
            )
            
            # Try to get token from cache first
            result = app.acquire_token_silent(self.scope, account=None)
            
            if not result:
                print("No suitable token in cache. Getting new token...", flush=True)
                result = app.acquire_token_for_client(scopes=self.scope)
            
            if result and 'access_token' in result:
                print("Successfully obtained access token")
                return result['access_token']
            else:
                error = result.get('error') if result else 'No result'
                error_desc = result.get('error_description') if result else 'No error description'
                print(f"Token acquisition failed: {error} - {error_desc}")
                raise Exception(f"Failed to get token: {error} - {error_desc}")
            
        except Exception as e:
            print(f"Error getting access token: {str(e)}", flush=True)
            raise

    def fetch_emails(self, days: int = None) -> list[dict[str, Any]]:
        """Fetch all emails from Microsoft Graph API (no date limit) or from the last N days if specified"""
        try:
            access_token = self._get_access_token()
            if not access_token:
                print("Failed to get access token")
                return []
            
            # Initialize variables for pagination - fetch only from inbox
            next_link = f"https://graph.microsoft.com/v1.0/users/{self.user_email}/mailFolders/inbox/messages"
            headers = {'Authorization': f'Bearer {access_token}'}
            page_size = 50
            all_emails = []
            
            while next_link:
                # Prepare parameters for the current page
                params = {
                    '$select': 'id,subject,body,receivedDateTime,sender',
                    '$orderby': 'receivedDateTime DESC',
                    '$top': page_size
                }
                
                # Add date filter only if days parameter is specified
                if days is not None:
                    end_date = datetime.utcnow()
                    start_date = end_date - timedelta(days=days)
                    start_date_str = start_date.strftime('%Y-%m-%dT%H:%M:%SZ')
                    end_date_str = end_date.strftime('%Y-%m-%dT%H:%M:%SZ')
                    params['$filter'] = f"receivedDateTime ge {start_date_str} and receivedDateTime le {end_date_str}"
                
                print(f"\nMaking request to: {next_link}", flush=True)
                print(f"With params: {params}", flush=True)
                
                response = requests.get(next_link, headers=headers, params=params)
                print(f"Response status: {response.status_code}", flush=True)
                
                if response.status_code != 200:
                    print(f"Error fetching emails: {response.text}")
                    break
                
                data = response.json()
                emails = data.get('value', [])
                
                for email in emails:
                    try:
                        # Process the email
                        processed_email = self._process_email(email)
                        
                        # Add to response
                        all_emails.append(processed_email)
                        
                    except Exception as e:
                        print(f"Error processing email {email.get('id')}: {str(e)}")
                        continue
                
                # Check for more pages
                next_link = data.get('@odata.nextLink')
                if not next_link:
                    break
            
            print(f"\nProcessed {len(all_emails)} emails")
            return all_emails
            
        except Exception as e:
            print(f"Error in fetch_emails: {str(e)}")
            return []

    def fetch_emails_since(self, since_datetime) -> list[dict[str, Any]]:
        """Fetch emails received after the specified datetime"""
        try:
            access_token = self._get_access_token()
            if not access_token:
                print("Failed to get access token")
                return []
            
            # Format the datetime for Microsoft Graph API
            since_str = since_datetime.strftime('%Y-%m-%dT%H:%M:%SZ')
            
            # Initialize variables for pagination - fetch only from inbox
            next_link = f"https://graph.microsoft.com/v1.0/users/{self.user_email}/mailFolders/inbox/messages"
            headers = {'Authorization': f'Bearer {access_token}'}
            page_size = 50
            all_emails = []
            
            while next_link:
                # Prepare parameters for the current page
                params = {
                    '$select': 'id,subject,body,receivedDateTime,sender',
                    '$orderby': 'receivedDateTime DESC',
                    '$top': page_size,
                    '$filter': f"receivedDateTime gt {since_str}"  # Only emails after last refresh
                }
                
                print(f"\nMaking request to: {next_link}", flush=True)
                print(f"With params: {params}", flush=True)
                
                response = requests.get(next_link, headers=headers, params=params)
                print(f"Response status: {response.status_code}", flush=True)
                
                if response.status_code != 200:
                    print(f"Error fetching emails: {response.text}")
                    break
                
                data = response.json()
                emails = data.get('value', [])
                
                for email in emails:
                    try:
                        # Process the email
                        processed_email = self._process_email(email)
                        
                        # Add to response
                        all_emails.append(processed_email)
                        
                    except Exception as e:
                        print(f"Error processing email {email.get('id')}: {str(e)}")
                        continue
                
                # Check for more pages
                next_link = data.get('@odata.nextLink')
                if not next_link:
                    break
            
            print(f"\nProcessed {len(all_emails)} emails since {since_str}")
            return all_emails
            
        except Exception as e:
            print(f"Error in fetch_emails_since: {str(e)}")
            return []